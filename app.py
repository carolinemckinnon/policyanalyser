import streamlit as st
from preprocessing import clean_text, remove_high_frequency_terms, basic_sentence_tokenize
from similarity import cluster_segments, label_clusters_by_tfidf
from collections import Counter
import pandas as pd
import numpy as np
from typing import Any, Dict, List, Optional, Tuple, Set
import os
import re
import hashlib
try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    Document = None
    PackageNotFoundError = Exception
os.environ.setdefault("TOKENIZERS_PARALLELISM", "false")

from io import BytesIO
from zipfile import ZipFile, BadZipFile
from pathlib import Path
LEGISLATION_SAMPLE_FILE = Path("legislation_texts/youth_justice_act_2024_sections.xlsx")
def load_document_from_register(entry: pd.Series, docs_dir: Optional[Path]) -> Dict:
    """Load policy text either from a DOCX file or fall back to the register body text."""
    metadata = entry.to_dict()
    display_name = entry.get("display_title") or entry.get("title") or entry.get("file_name")
    body_text = str(entry.get("body_text") or "").strip()
    raw_text = body_text
    paragraphs: List = []
    source = metadata.get("relative_path", "")

    if docs_dir is not None and metadata.get("relative_path"):
        doc_path = (docs_dir / metadata["relative_path"]).expanduser().resolve()
        if Document is not None and doc_path.exists():
            try:
                document = Document(str(doc_path))
            except Exception:  # pragma: no cover - fallback to register text
                document = None
            if document is not None:
                paragraphs = [p for p in document.paragraphs if p.text.strip()]
                raw_text = "\n".join(p.text.strip() for p in paragraphs)
                source = str(doc_path)

    if not raw_text:
        raw_text = body_text

    return {
        "name": display_name,
        "raw_text": raw_text,
        "paragraphs": paragraphs,
        "source": source,
        "register_id": int(entry["register_id"]),
        "metadata": metadata,
    }



import subprocess
import sys
from urllib.parse import quote, unquote

BASE_DIR = Path(__file__).resolve().parent
DEFAULT_DATA_DIR = BASE_DIR / "data"
DEFAULT_REGISTER_FILENAME = "document_register_enriched.csv"
REGISTER_EMBEDDINGS_FILENAME = "document_register_enriched_embeddings.npz"
CHUNK_METADATA_PATH = DEFAULT_DATA_DIR / "semantic_space" / "policy_chunks.csv"
CHUNK_EMBEDDINGS_PATH = DEFAULT_DATA_DIR / "semantic_space" / "policy_chunks_embeddings.npy"
DEFAULT_DOC_DIR_CANDIDATES = [
    DEFAULT_DATA_DIR / "docs",
    Path("./data/docs"),
    Path("../data/docs"),
]


@st.cache_data(show_spinner=False)
def load_register_data(register_path: str) -> pd.DataFrame:
    df = pd.read_csv(register_path)
    df["register_id"] = df["register_id"].astype(int)
    for col in ("first_published", "last_revised"):
        if col in df.columns:
            df[col] = pd.to_datetime(df[col], errors="coerce", dayfirst=True)
    if "title" in df.columns:
        df["title"] = df["title"].fillna("")
    if "file_name" in df.columns:
        df["file_name"] = df["file_name"].fillna("")
    df["display_title"] = df["title"].where(df["title"].str.strip() != "", df["file_name"])
    if "last_revised" in df.columns:
        df["last_revised"] = pd.to_datetime(df["last_revised"], errors="coerce", dayfirst=True)
    return df


def guess_default_docs_dir() -> str:
    for candidate in DEFAULT_DOC_DIR_CANDIDATES:
        candidate_path = candidate.expanduser().resolve()
        if candidate_path.exists():
            register_path = candidate_path / DEFAULT_REGISTER_FILENAME
            if register_path.exists():
                return str(candidate_path)
    return ""


@st.cache_data(show_spinner=False)
def load_ontology_data(base_dir: str) -> Optional[pd.DataFrame]:
    path = Path(base_dir) / "policy_ontology.csv"
    if not path.exists():
        return None
    df = pd.read_csv(path)
    df["register_id"] = df["register_id"].astype(int)
    df["legislation_reference"] = df["legislation_reference"].fillna("")
    return df


@st.cache_data(show_spinner=False)
def load_definition_conflicts(base_dir: str) -> Optional[pd.DataFrame]:
    path = Path(base_dir) / "definition_conflicts.csv"
    if not path.exists():
        return None
    return pd.read_csv(path)


@st.cache_resource(show_spinner=False)
def load_sentence_transformer():
    from sentence_transformers import SentenceTransformer
    return SentenceTransformer("all-MiniLM-L6-v2")


@st.cache_data(show_spinner=False)
def compute_text_embeddings(texts_tuple):
    if not texts_tuple:
        return np.empty((0, 0))
    model = load_sentence_transformer()
    return model.encode(list(texts_tuple), normalize_embeddings=True)


def compute_cosine_similarity_matrix(vectors: np.ndarray) -> np.ndarray:
    if vectors is None or getattr(vectors, "size", 0) == 0:
        return np.empty((0, 0))
    norms = np.linalg.norm(vectors, axis=1, keepdims=True)
    norms = np.where(norms == 0, 1, norms)
    normalized = vectors / norms
    return normalized @ normalized.T
@st.cache_data(show_spinner=False)
def compute_register_embeddings(ids_tuple, texts_tuple):
    if not ids_tuple:
        return np.empty((0, 0))
    return compute_text_embeddings(texts_tuple)


def build_legislation_groups(df: pd.DataFrame) -> dict:
    groups = {
        'Section': [],
        'Part': [],
        'Chapter': [],
    }
    for _, row in df.iterrows():
        context = row['legislation_section']
        content = row['content']
        segment_number = row['segment_number']
        groups['Section'].append({
            'label': context,
            'content': content,
            'segment_numbers': [segment_number],
        })
        parts = context.split(' | ')
        part_label = None
        chapter_label = None
        for part in parts:
            lower = part.lower()
            if lower.startswith('part '):
                part_label = part
            if lower.startswith('chapter '):
                chapter_label = part
        if part_label:
            entry = next((g for g in groups['Part'] if g['label'] == part_label), None)
            if entry is None:
                entry = {'label': part_label, 'content_list': [], 'segment_numbers': []}
                groups['Part'].append(entry)
            entry['content_list'].append(content)
            entry['segment_numbers'].append(segment_number)
        if chapter_label:
            entry = next((g for g in groups['Chapter'] if g['label'] == chapter_label), None)
            if entry is None:
                entry = {'label': chapter_label, 'content_list': [], 'segment_numbers': []}
                groups['Chapter'].append(entry)
            entry['content_list'].append(content)
            entry['segment_numbers'].append(segment_number)
    for level in ('Part', 'Chapter'):
        for entry in groups[level]:
            entry['content'] = ' '.join(entry.pop('content_list'))
    return groups


@st.cache_data(show_spinner=False)
def load_legislation_sections(path: str):
    file_path = Path(path)
    if not file_path.exists():
        return None, None, None
    df = pd.read_excel(file_path)
    df = df.fillna('')
    df['legislation_section'] = df['legislation_section'].astype(str)
    df['content'] = df['content'].astype(str)
    texts_tuple = tuple(df['content'].tolist())
    embeddings = compute_text_embeddings(texts_tuple)
    groups = build_legislation_groups(df)
    return df, embeddings, groups


def build_legislation_aliases(name: str) -> Set[str]:
    """Generate lowercase aliases of an Act name for matching policy references."""
    if not name:
        return set()
    cleaned = re.sub(r"\s+", " ", name).strip()
    cleaned = cleaned.replace("(sample)", "").strip()
    aliases = {cleaned.lower()}
    aliases.add(cleaned.lower().split("(")[0].strip())
    if cleaned.lower().endswith(" act"):
        aliases.add(cleaned.lower()[:-4].strip())
    if cleaned.lower().endswith(" act") and len(cleaned.split()) >= 2:
        aliases.add(" ".join(cleaned.lower().split()[:-1]))
    return {alias for alias in aliases if alias}


def build_policy_chunks(register_df: pd.DataFrame, sentences_per_chunk: int = 5) -> Tuple[pd.DataFrame, np.ndarray]:
    """Create paragraph/sentence chunks for each policy and compute embeddings."""
    records: List[Dict] = []
    for _, row in register_df.iterrows():
        register_id = int(row["register_id"])
        title = row.get("display_title") or row.get("title") or row.get("file_name") or f"Policy {register_id}"
        body_text = str(row.get("body_text") or "").strip()
        if not body_text:
            continue
        sentences = basic_sentence_tokenize(body_text)
        if not sentences:
            continue
        chunk_index = 0
        for start in range(0, len(sentences), sentences_per_chunk):
            chunk_sentences = sentences[start:start + sentences_per_chunk]
            if not chunk_sentences:
                continue
            chunk_text = " ".join(chunk_sentences).strip()
            if not chunk_text:
                continue
            records.append(
                {
                    "chunk_id": f"{register_id}_{chunk_index}",
                    "register_id": register_id,
                    "policy_title": title,
                    "chunk_index": chunk_index,
                    "start_sentence": start,
                    "end_sentence": start + len(chunk_sentences) - 1,
                    "chunk_text": chunk_text,
                    "chunk_word_count": len(chunk_text.split()),
                    "policy_word_count": row.get("word_count", 0),
                    "type": row.get("type", ""),
                    "policy_owner": row.get("policy_owner", ""),
                    "relative_path": row.get("relative_path", ""),
                    "policy_subject": row.get("policy_subject", ""),
                    "policy_sub_topic": row.get("policy_sub_topic", ""),
                    "executive_owner": row.get("executive_owner", ""),
                    "last_revised": (
                        row["last_revised"].date().isoformat()
                        if isinstance(row.get("last_revised"), pd.Timestamp) and not pd.isna(row.get("last_revised"))
                        else str(row.get("last_revised") or "")
                    ),
                }
            )
            chunk_index += 1

    chunk_df = pd.DataFrame(records).reset_index(drop=True)
    if chunk_df.empty:
        return chunk_df, np.empty((0, 0))

    embeddings = compute_text_embeddings(tuple(chunk_df["chunk_text"].tolist()))

    CHUNK_METADATA_PATH.parent.mkdir(parents=True, exist_ok=True)
    chunk_df.to_csv(CHUNK_METADATA_PATH, index=False)
    np.save(CHUNK_EMBEDDINGS_PATH, embeddings)
    return chunk_df, embeddings


def load_policy_chunks(register_df: pd.DataFrame) -> Tuple[pd.DataFrame, np.ndarray]:
    """Load chunk metadata and embeddings, rebuilding if necessary."""
    required_columns = {
        "chunk_id",
        "register_id",
        "policy_title",
        "chunk_index",
        "chunk_text",
        "chunk_word_count",
        "policy_word_count",
        "type",
        "policy_owner",
        "relative_path",
        "last_revised",
        "policy_subject",
        "policy_sub_topic",
        "executive_owner",
    }

    def needs_rebuild(df: pd.DataFrame, embeddings: np.ndarray) -> bool:
        if df.empty:
            return True
        if len(df) != embeddings.shape[0]:
            return True
        if not required_columns.issubset(set(df.columns)):
            return True
        cached_ids = set(df["register_id"].astype(int).unique())
        current_ids = set(register_df["register_id"].astype(int).unique())
        if cached_ids != current_ids:
            return True
        return False

    if CHUNK_METADATA_PATH.exists() and CHUNK_EMBEDDINGS_PATH.exists():
        chunk_df = pd.read_csv(CHUNK_METADATA_PATH).reset_index(drop=True)
        embeddings = np.load(CHUNK_EMBEDDINGS_PATH)
        if needs_rebuild(chunk_df, embeddings):
            chunk_df, embeddings = build_policy_chunks(register_df)
    else:
        chunk_df, embeddings = build_policy_chunks(register_df)

    return chunk_df, embeddings


def retrieve_relevant_chunks(
    query: str,
    chunk_df: pd.DataFrame,
    chunk_embeddings: np.ndarray,
    top_k: int = 5,
) -> pd.DataFrame:
    """Return the top chunk matches for a natural-language query."""
    if not query.strip() or chunk_embeddings.size == 0 or chunk_df.empty:
        return pd.DataFrame()
    model = load_sentence_transformer()
    query_vec = model.encode([query], normalize_embeddings=True)[0]
    scores = chunk_embeddings @ query_vec
    result_df = chunk_df.copy()
    result_df["score"] = scores
    return result_df.sort_values("score", ascending=False).head(top_k)


def filter_chunk_subset(
    chunk_df: pd.DataFrame,
    chunk_embeddings: np.ndarray,
    active_ids: Set[int],
    selected_subjects: List[str],
    selected_subtopics: List[str],
    selected_owners: List[str],
    refine_terms: List[str],
) -> Tuple[pd.DataFrame, np.ndarray]:
    """Apply filters to the chunk table and return the subset with aligned embeddings."""
    if chunk_df.empty or chunk_embeddings.size == 0 or not active_ids:
        return pd.DataFrame(), np.empty((0, 0))

    available = chunk_df[chunk_df["register_id"].isin(active_ids)].copy()
    if selected_subjects:
        available = available[available["policy_subject"].fillna("").isin(selected_subjects)]
    if selected_subtopics:
        available = available[available["policy_sub_topic"].fillna("").isin(selected_subtopics)]
    if selected_owners:
        available = available[available["policy_owner"].fillna("").isin(selected_owners)]

    if refine_terms:
        mask = pd.Series(True, index=available.index)
        for term in refine_terms:
            mask &= available["chunk_text"].str.contains(term, case=False, na=False)
        available = available[mask]

    if available.empty:
        return pd.DataFrame(), np.empty((0, 0))

    indices = available.index.to_numpy(dtype=int)
    embeddings_subset = chunk_embeddings[indices]
    available = available.reset_index(drop=True)
    return available, embeddings_subset


def answer_question_extractive(
    question: str,
    available_chunks: pd.DataFrame,
    embeddings_subset: np.ndarray,
    docs_dir: Optional[Path],
    top_k: int = 5,
) -> List[Dict]:
    """Return extractive answers drawn from the highest-scoring chunks."""
    if not question.strip() or available_chunks.empty or embeddings_subset.size == 0:
        return []

    model = load_sentence_transformer()
    query_vec = model.encode([question], normalize_embeddings=True)[0]
    scores = embeddings_subset @ query_vec

    candidates = available_chunks.copy()
    candidates["score"] = scores
    top_chunks = candidates.sort_values("score", ascending=False).head(top_k)

    results: List[Dict] = []
    for _, chunk in top_chunks.iterrows():
        sentences = basic_sentence_tokenize(chunk["chunk_text"])
        best_sentence = chunk["chunk_text"].strip()
        support_sentences: List[str] = []

        if sentences:
            sentence_embeddings = compute_text_embeddings(tuple(sentences))
            sentence_scores = sentence_embeddings @ query_vec
            best_idx = int(np.argmax(sentence_scores))
            best_sentence = sentences[best_idx].strip()
            supporting_indices = [idx for idx in np.argsort(sentence_scores)[::-1] if idx != best_idx]
            for idx in supporting_indices[:2]:
                sent = sentences[idx].strip()
                if sent and sent not in support_sentences:
                    support_sentences.append(sent)

        document_link = ""
        rel_path = chunk.get("relative_path")
        if docs_dir and rel_path:
            doc_path = (Path(docs_dir) / rel_path).expanduser().resolve()
            if doc_path.exists():
                document_link = f"?open_file={quote(str(doc_path))}"

        results.append(
            {
                "policy": chunk.get("policy_title"),
                "register_id": int(chunk.get("register_id", 0)),
                "score": float(chunk.get("score", 0.0)),
                "policy_subject": chunk.get("policy_subject"),
                "policy_sub_topic": chunk.get("policy_sub_topic"),
                "policy_owner": chunk.get("policy_owner"),
                "best_sentence": best_sentence,
                "support_sentences": support_sentences,
                "document_link": document_link,
            }
        )

    return results


def classify_legislation_heading(text: str, style_name: str) -> Tuple[Optional[str], Optional[str]]:
    """Identify whether a paragraph represents a chapter, part, division, or section heading."""
    cleaned = text.strip()
    if not cleaned:
        return None, None
    style_lower = (style_name or "").lower()
    normalized = re.sub(r"\s+", " ", cleaned)
    lower = normalized.lower()

    chapter_match = re.match(r"^(chapter\s+\d+[a-z]?\b.*)", lower)
    if chapter_match:
        return "chapter", normalized

    part_match = re.match(r"^(part\s+\d+[a-z]?\b.*)", lower)
    if part_match:
        return "part", normalized

    division_match = re.match(r"^(division\s+\d+[a-z]?\b.*)", lower)
    if division_match:
        return "division", normalized

    if re.match(r"^(schedule\s+\d+[a-z]?\b.*)", lower):
        return "chapter", normalized

    if re.match(r"^(section\s+\d+[a-z]?\b.*)", lower):
        return "section", normalized

    heading_style = any(
        token in style_lower
        for token in ("heading", "title", "section heading", "chapter heading")
    ) or style_lower.startswith("h")

    if heading_style and re.match(r"^\d+[a-z]?\s+[A-Z]", normalized) and len(normalized.split()) <= 20:
        return "section", normalized

    if heading_style and re.match(r"^\d+[a-z]?\.\s+[A-Z]", normalized) and len(normalized.split()) <= 20:
        return "section", normalized

    return None, None


def extract_legislation_title(document: Document) -> Optional[str]:
    """Attempt to pull a suitable title from the document."""
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if text.isupper() and "ACT" in text:
            return text
        if " Act" in text and len(text) < 120:
            return text
    return None


def process_legislation_bytes(file_bytes: bytes, file_name: str) -> Dict:
    """Parse a DOCX Act into sections, compute embeddings, and prepare grouping metadata."""
    document = Document(BytesIO(file_bytes))
    current_chapter: Optional[str] = None
    current_part: Optional[str] = None
    current_division: Optional[str] = None
    current_section_heading: Optional[str] = None
    current_lines: List[str] = []
    preface_lines: List[str] = []
    rows: List[Dict[str, str]] = []
    segment_number = 0
    started_content = False

    def flush_section():
        nonlocal current_section_heading, current_lines, segment_number
        if not current_lines:
            return
        segment_number += 1
        context_parts = [
            part
            for part in [
                current_chapter,
                current_part,
                current_division,
                current_section_heading,
            ]
            if part
        ]
        label = " | ".join(context_parts) if context_parts else f"Section {segment_number}"
        content_text = normalize_policy_text(" ".join(current_lines))
        rows.append(
            {
                "segment_number": segment_number,
                "legislation_section": label,
                "content": content_text,
            }
        )
        current_lines = []
        current_section_heading = None

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = getattr(para.style, "name", "") or ""
        heading_type, normalized_heading = classify_legislation_heading(text, style_name)

        lower_text = text.lower()
        if not started_content:
            final_token = text.split()[-1]
            looks_like_page = final_token.isdigit()
            if "table of contents" in lower_text or lower_text in {"contents", "table of provisions"}:
                continue
            if looks_like_page or "...." in text or "\t" in text:
                # Appears to be a table-of-contents entry with page numbers or dot leaders.
                continue
            if heading_type:
                started_content = True
            else:
                # Skip any other front-matter paragraphs.
                continue

        if heading_type == "chapter":
            current_chapter = normalized_heading
            current_part = None
            current_division = None
            continue

        if heading_type == "part":
            current_part = normalized_heading
            current_division = None
            continue

        if heading_type == "division":
            current_division = normalized_heading
            continue

        if heading_type == "section":
            if current_lines:
                flush_section()
            current_section_heading = normalized_heading
            current_lines = [normalized_heading]
            continue

        if current_section_heading is None:
            preface_lines.append(text)
        else:
            current_lines.append(text)

    if current_lines:
        flush_section()

    if rows and preface_lines:
        rows[0]["content"] = normalize_policy_text(" ".join(preface_lines + [rows[0]["content"]]))
    elif preface_lines:
        rows.insert(
            0,
            {
                "segment_number": 1,
                "legislation_section": "Introductory",
                "content": normalize_policy_text(" ".join(preface_lines)),
            },
        )
        for idx, row in enumerate(rows, start=1):
            row["segment_number"] = idx

    if not rows:
        raise ValueError("No sections were detected in the uploaded document. Check the formatting and try again.")

    df = pd.DataFrame(rows)
    df = df.fillna("")
    df["segment_number"] = range(1, len(df) + 1)
    df["legislation_section"] = df["legislation_section"].astype(str)
    df["content"] = df["content"].astype(str)

    embeddings = compute_text_embeddings(tuple(df["content"].tolist()))
    groups = build_legislation_groups(df)
    title = extract_legislation_title(document) or Path(file_name).stem.replace("_", " ")
    aliases = build_legislation_aliases(title)
    cache_key = hashlib.sha256(file_bytes).hexdigest()[:16]
    return {
        "df": df,
        "embeddings": embeddings,
        "groups": groups,
        "display_name": title,
        "source_name": file_name,
        "cache_key": cache_key,
        "section_count": len(df),
        "aliases": aliases,
    }


def normalize_policy_text(text: str) -> str:
    return re.sub(r'\s+', ' ', text).strip()


def ensure_timestamp_naive(value: Any) -> Optional[pd.Timestamp]:
    """Convert various timestamp inputs to timezone-naive pandas.Timestamp."""
    if value is None or (isinstance(value, (float, np.floating)) and np.isnan(value)):
        return None
    if isinstance(value, pd.Timestamp):
        ts = value
    else:
        ts = pd.to_datetime(value, errors="coerce")
    if ts is None or pd.isna(ts):
        return None
    tzinfo = getattr(ts, "tzinfo", None)
    if tzinfo is not None:
        try:
            ts = ts.tz_localize(None)
        except TypeError:
            ts = ts.tz_convert(None).tz_localize(None)
    return ts


def normalize_text_field(value: Any) -> str:
    """Return a safe, trimmed string for metadata values (handles NaN, None, etc.)."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (float, np.floating)) and np.isnan(value):
        return ""
    return str(value).strip()


def ensure_comparison_styles():
    """Inject CSS styling for policy comparison blocks."""
    st.markdown(
        """
        <style>
            .compare-block {
                background: #f7f9fc;
                border: 1px solid #dfe3eb;
                border-radius: 8px;
                padding: 12px;
                margin-bottom: 10px;
                font-size: 0.95rem;
            }
            .compare-title {
                font-weight: 600;
                margin-bottom: 6px;
            }
            .match-highlight {
                background-color: #ffe9a8;
                border-radius: 3px;
                padding: 0 2px;
            }
            .unique-a {
                background-color: #d9f0ff;
                border-radius: 3px;
                padding: 0 2px;
            }
            .unique-b {
                background-color: #e8d9ff;
                border-radius: 3px;
                padding: 0 2px;
            }
        </style>
        """,
        unsafe_allow_html=True,
    )


def build_highlighted_text(
    text: str,
    overlap_tokens: Set[str],
    unique_tokens: Set[str],
    unique_class: str,
) -> str:
    """Highlight overlapping and unique tokens within text."""
    cleaned = re.sub(r"\s+", " ", text.strip())
    if not cleaned:
        return ""
    token_map: Dict[str, str] = {}
    for token in overlap_tokens:
        if token:
            token_map[token.lower()] = "match-highlight"
    for token in unique_tokens:
        if token and token.lower() not in token_map:
            token_map[token.lower()] = unique_class
    if not token_map:
        return re.sub(r"\s+", " ", text.strip())

    escaped_tokens = sorted({re.escape(token) for token in token_map.keys()}, key=len, reverse=True)
    pattern = re.compile(r"\b(" + "|".join(escaped_tokens) + r")\b", flags=re.IGNORECASE)

    def repl(match: re.Match) -> str:
        token_lower = match.group(0).lower()
        css_class = token_map.get(token_lower, "")
        return f'<span class="{css_class}">{match.group(0)}</span>' if css_class else match.group(0)

    return pattern.sub(repl, cleaned)


def extract_top_matching_snippets(
    text_a: str,
    text_b: str,
    top_n: int = 3,
    max_sentences: int = 200,
) -> Tuple[
    List[Tuple[str, str, float]],
    List[Tuple[str, str, float]],
    List[Tuple[str, str, float]],
    List[Tuple[str, str, float]],
]:
    """Return the strongest matching and most divergent snippets between two policy texts."""
    sentences_a = [
        sent.strip()
        for sent in basic_sentence_tokenize(text_a)
        if isinstance(sent, str) and len(sent.strip()) >= 40
    ][:max_sentences]
    sentences_b = [
        sent.strip()
        for sent in basic_sentence_tokenize(text_b)
        if isinstance(sent, str) and len(sent.strip()) >= 40
    ][:max_sentences]

    if not sentences_a:
        sentences_a = [text_a.strip()] if text_a.strip() else []
    if not sentences_b:
        sentences_b = [text_b.strip()] if text_b.strip() else []
    if not sentences_a or not sentences_b:
        return [], [], [], []

    emb_a = compute_text_embeddings(tuple(sentences_a))
    emb_b = compute_text_embeddings(tuple(sentences_b))
    if emb_a.size == 0 or emb_b.size == 0:
        return [], [], [], []

    sim_matrix = emb_a @ emb_b.T
    if sim_matrix.size == 0:
        return [], [], [], []

    best_matches_a: List[Tuple[int, int, float]] = []
    for idx in range(sim_matrix.shape[0]):
        best_col = int(np.argmax(sim_matrix[idx]))
        best_matches_a.append((idx, best_col, float(sim_matrix[idx, best_col])))
    best_matches_a.sort(key=lambda item: item[2], reverse=True)

    selected_a: List[Tuple[str, str, float]] = []
    used_indices_a: Set[int] = set()
    used_indices_b: Set[int] = set()
    for idx, col, score in best_matches_a:
        if len(selected_a) >= top_n:
            break
        if idx in used_indices_a or col in used_indices_b:
            continue
        selected_a.append((sentences_a[idx], sentences_b[col], score))
        used_indices_a.add(idx)
        used_indices_b.add(col)

    best_matches_b: List[Tuple[int, int, float]] = []
    for col in range(sim_matrix.shape[1]):
        best_row = int(np.argmax(sim_matrix[:, col]))
        best_matches_b.append((col, best_row, float(sim_matrix[best_row, col])))
    best_matches_b.sort(key=lambda item: item[2], reverse=True)

    selected_b: List[Tuple[str, str, float]] = []
    used_indices_b_set: Set[int] = set()
    used_indices_a_for_b: Set[int] = set()
    for col, row, score in best_matches_b:
        if len(selected_b) >= top_n:
            break
        if col in used_indices_b_set or row in used_indices_a_for_b:
            continue
        selected_b.append((sentences_b[col], sentences_a[row], score))
        used_indices_b_set.add(col)
        used_indices_a_for_b.add(row)

    # Divergence: sentences with the lowest maximum similarity.
    max_scores_a = sim_matrix.max(axis=1)
    best_cols_a = sim_matrix.argmax(axis=1)
    divergence_candidates_a = sorted(
        [(idx, int(best_cols_a[idx]), float(max_scores_a[idx])) for idx in range(len(sentences_a))],
        key=lambda item: item[2],
    )
    divergence_a: List[Tuple[str, str, float]] = []
    seen_a: Set[int] = set()
    for idx, col, score in divergence_candidates_a:
        if len(divergence_a) >= top_n:
            break
        if idx in seen_a or (idx in used_indices_a and col in used_indices_b):
            continue
        divergence_a.append((sentences_a[idx], sentences_b[col], score))
        seen_a.add(idx)

    max_scores_b = sim_matrix.max(axis=0)
    best_rows_b = sim_matrix.argmax(axis=0)
    divergence_candidates_b = sorted(
        [(col, int(best_rows_b[col]), float(max_scores_b[col])) for col in range(len(sentences_b))],
        key=lambda item: item[2],
    )
    divergence_b: List[Tuple[str, str, float]] = []
    seen_b: Set[int] = set()
    for col, row, score in divergence_candidates_b:
        if len(divergence_b) >= top_n:
            break
        if col in seen_b or (row in used_indices_a and col in used_indices_b):
            continue
        divergence_b.append((sentences_b[col], sentences_a[row], score))
        seen_b.add(col)

    return selected_a, selected_b, divergence_a, divergence_b


def build_pair_summary(
    label_a: str,
    label_b: str,
    detail: Dict,
    subject_a: str,
    subject_b: str,
    subtopic_a: str,
    subtopic_b: str,
    subject_match: bool,
    subtopic_match: bool,
    matches_a: List[Tuple[str, str, float]],
    matches_b: List[Tuple[str, str, float]],
    unique_tokens_a: Set[str],
    unique_tokens_b: Set[str],
) -> str:
    """Create a brief natural-language summary for a policy pair."""
    segments: List[str] = []
    hybrid_score = detail.get("hybrid_score")
    if hybrid_score is not None:
        segments.append(
            f"{label_a} and {label_b} score {hybrid_score:.2f} on the analyser, indicating strong overlap."
        )

    if subject_match and subject_a:
        segments.append(f"Both policies sit under the **{subject_a}** subject.")
    elif subject_a or subject_b:
        human_subject_a = subject_a or "unspecified subject"
        human_subject_b = subject_b or "unspecified subject"
        segments.append(
            f"They are filed under **{human_subject_a}** and **{human_subject_b}** respectively, so check the fit before merging."
        )

    if subtopic_match and subtopic_a:
        segments.append(f"They also share the **{subtopic_a}** sub-topic.")

    shared_terms = detail.get("shared_terms") or []
    if shared_terms:
        preview_terms = ", ".join(shared_terms[:3])
        if len(shared_terms) > 3:
            preview_terms += ", …"
        segments.append(f"Common language includes {preview_terms}.")

    shared_legislation = detail.get("shared_legislation") or []
    if shared_legislation:
        preview_acts = ", ".join(shared_legislation[:2])
        if len(shared_legislation) > 2:
            preview_acts += ", …"
        segments.append(f"Both reference {preview_acts}.")

    conflicting_terms = detail.get("conflicting_terms") or []
    if conflicting_terms:
        preview_conflicts = ", ".join(conflicting_terms[:2])
        if len(conflicting_terms) > 2:
            preview_conflicts += ", …"
        segments.append(f"Watch for definition differences on {preview_conflicts}.")

    def shortlist_tokens(tokens: Set[str]) -> Optional[str]:
        if not tokens:
            return None
        selected = sorted(tokens)
        preview = ", ".join(selected[:3])
        if len(selected) > 3:
            preview += ", …"
        return preview

    return " ".join(segments)


def clear_cached_results():
    for key in [
        "document_similarity_matrix",
        "text_similarity_matrix",
        "document_labels",
        "document_snippets",
        "filtered_terms",
        "document_cluster_labels",
        "document_theme_keywords",
        "document_embeddings",
        "document_ids",
        "document_texts_tuple",
        "loaded_documents",
        "load_errors",
        "pairwise_similarity_details",
        "document_terms",
        "document_legislation",
        "hybrid_parameters",
    ]:
        st.session_state.pop(key, None)


def load_documents_from_ids(
    selected_ids: List[str],
    register_df: pd.DataFrame,
    docs_dir: Path,
) -> Tuple[List[Dict], List[str]]:
    documents = []
    errors: List[str] = []
    for sid in selected_ids:
        try:
            rid = int(sid)
        except (TypeError, ValueError):
            errors.append(f"Invalid register ID: {sid}")
            continue
        entry = register_df.loc[register_df["register_id"] == rid]
        if entry.empty:
            errors.append(f"Register entry {rid} not found.")
            continue
        entry_series = entry.iloc[0]
        try:
            record = load_document_from_register(entry_series, docs_dir)
        except Exception as exc:
            errors.append(f"{rid}: {exc}")
            continue
        documents.append(record)
    return documents, errors


def compute_hybrid_scores(
    base_vector: np.ndarray,
    doc_embeddings: np.ndarray,
    doc_ids: List[Optional[int]],
    term_map: Dict[int, Dict[str, str]],
    legislation_map: Dict[int, Set[str]],
    ontology_weight: float,
    term_bias: float,
    section_text_lower: str,
    act_aliases: Set[str],
) -> List[Dict]:
    """Compute hybrid similarity scores between a base vector and cached policy embeddings."""
    if doc_embeddings is None or doc_embeddings.size == 0:
        return []

    text_scores = doc_embeddings @ base_vector
    term_weight = max(term_bias, 0.0)
    legislation_weight = max(1.0 - term_bias, 0.0)
    if term_weight + legislation_weight == 0:
        term_weight = legislation_weight = 0.5

    results = []
    for idx, base_score in enumerate(text_scores):
        rid = doc_ids[idx] if idx < len(doc_ids) else None
        raw_terms = term_map.get(rid) if isinstance(rid, int) else None
        raw_legs = legislation_map.get(rid) if isinstance(rid, int) else None

        if isinstance(raw_terms, dict):
            terms_i = raw_terms
        elif isinstance(raw_terms, (list, set, tuple)):
            terms_i = {term: "" for term in raw_terms}
        else:
            terms_i = {}

        if isinstance(raw_legs, set):
            legs_i = raw_legs
        elif isinstance(raw_legs, (list, tuple)):
            legs_i = set(raw_legs)
        else:
            legs_i = set()

        term_union: Set[str] = set()
        shared_terms: Set[str] = set()
        align_terms: Set[str] = set()
        conflict_terms: Set[str] = set()
        legislation_union: Set[str] = set()
        shared_legislation: Set[str] = set()

        if terms_i:
            term_union = set(terms_i.keys())
            for term in term_union:
                if term.lower() in section_text_lower:
                    shared_terms.add(term)
            align_terms = set(shared_terms)

        if legs_i:
            legislation_union = set(legs_i)
            for ref in legs_i:
                ref_lower = ref.lower()
                if any(alias and alias in ref_lower for alias in act_aliases):
                    shared_legislation.add(ref)

        term_score = len(shared_terms) / len(term_union) if term_union else 0.0
        legislation_score = len(shared_legislation) / len(legislation_union) if legislation_union else 0.0

        ontology_component = 0.0
        if term_score > 0 or legislation_score > 0:
            numerator = term_weight * term_score + legislation_weight * legislation_score
            ontology_component = numerator / (term_weight + legislation_weight)

        hybrid_score = ((1 - ontology_weight) * float(base_score)) + (ontology_weight * ontology_component)
        hybrid_score = max(0.0, min(1.0, hybrid_score))

        results.append(
            {
                "index": idx,
                "hybrid_score": hybrid_score,
                "text_score": float(base_score),
                "term_score": term_score,
                "legislation_score": legislation_score,
                "shared_terms": sorted(shared_terms),
                "aligned_terms": sorted(align_terms),
                "conflicting_terms": sorted(conflict_terms),
                "shared_legislation": sorted(shared_legislation),
                "all_legislation": sorted(legislation_union),
                "all_terms": sorted(term_union),
            }
        )
    return results


def ensure_policy_data(filtered_df: pd.DataFrame, register_df: pd.DataFrame, docs_dir: Optional[Path]) -> bool:
    """Ensure policy documents and embeddings are ready for legislation matching."""
    documents = st.session_state.get("loaded_documents")
    doc_embeddings = st.session_state.get("document_embeddings")
    doc_ids = st.session_state.get("document_ids")
    doc_texts_tuple = st.session_state.get("document_texts_tuple")

    if documents and doc_embeddings is not None and getattr(doc_embeddings, "size", 0) > 0 and doc_ids:
        return True

    if docs_dir is None:
        st.warning("Set a valid documents directory to compare policies against legislation.")
        return False

    doc_tokens = filtered_df["register_id"].astype(str).tolist()
    if not doc_tokens:
        st.info("No policies match the current filters.")
        return False

    with st.spinner("Preparing policy data…"):
        documents, errors = load_documents_from_ids(doc_tokens, register_df, docs_dir)
        if errors:
            preview = "\n".join(errors[:5])
            st.warning(f"Some policies could not be loaded:\n{preview}")
        if not documents:
            st.info("No policies available to compare. Check the filters or document directory.")
            return False

        doc_labels: List[str] = []
        doc_texts: List[str] = []
        doc_id_list: List[Optional[int]] = []
        for doc in documents:
            label = doc.get("name") or doc.get("title") or doc.get("source") or "Document"
            doc_labels.append(label)
            text = normalize_policy_text(doc.get("raw_text", ""))
            doc_texts.append(text)
            try:
                doc_id_list.append(int(doc.get("register_id")))
            except (TypeError, ValueError):
                doc_id_list.append(None)

        embeddings = compute_text_embeddings(tuple(doc_texts)) if doc_texts else np.empty((0, 0))
        if embeddings.size == 0:
            st.info("Unable to compute policy embeddings for the current selection.")
            return False

        term_map: Dict[int, Dict[str, str]] = {}
        legislation_map: Dict[int, Set[str]] = {}
        ontology_df = load_ontology_data(str(docs_dir))
        if ontology_df is not None:
            ontology_df = ontology_df.fillna("")
            scoped = ontology_df[ontology_df["register_id"].isin(
                [rid for rid in doc_id_list if isinstance(rid, int)]
            )]
            for rid, group in scoped.groupby("register_id"):
                term_map[rid] = {
                    row.term: normalize_policy_text(row.definition)
                    for row in group.itertuples()
                    if row.term
                }
                leg_set: Set[str] = set()
                for ref in group["legislation_reference"]:
                    if not ref:
                        continue
                    parts = [part.strip() for part in ref.split(";")]
                    leg_set.update([part for part in parts if part])
                legislation_map[rid] = leg_set

        st.session_state["loaded_documents"] = documents
        st.session_state["document_embeddings"] = embeddings
        st.session_state["document_ids"] = doc_id_list
        st.session_state["document_texts_tuple"] = tuple(doc_texts)
        st.session_state["document_labels"] = doc_labels
        st.session_state["document_terms"] = term_map
        st.session_state["document_legislation"] = legislation_map
        st.session_state.pop("document_similarity_matrix", None)
        st.session_state.pop("text_similarity_matrix", None)
        st.session_state.pop("pairwise_similarity_details", None)

    return True


def load_or_build_register_embeddings(register_df: pd.DataFrame, docs_dir: Path) -> Tuple[np.ndarray, np.ndarray]:
    embedding_path = docs_dir / REGISTER_EMBEDDINGS_FILENAME
    register_ids = register_df["register_id"].to_numpy(dtype=np.int64)
    if embedding_path.exists():
        try:
            with np.load(embedding_path) as data:
                embeddings = data["embeddings"]
                stored_ids = data["register_ids"]
            if embeddings.shape[0] == len(register_df) and np.array_equal(stored_ids, register_ids):
                return embeddings, register_ids
        except Exception as exc:  # pylint: disable=broad-except
            st.warning(f"Could not read cached register embeddings ({embedding_path.name}): {exc}")

    normalized_texts = register_df.get("body_text", pd.Series([""] * len(register_df)))
    normalized_texts = normalized_texts.fillna("").apply(normalize_policy_text)
    embeddings = compute_text_embeddings(tuple(normalized_texts))
    try:
        np.savez_compressed(embedding_path, embeddings=embeddings, register_ids=register_ids)
    except Exception as exc:  # pylint: disable=broad-except
        st.warning(f"Could not write register embeddings cache ({embedding_path.name}): {exc}")
    return embeddings, register_ids


def render_document_similarity_tab(
    filtered_df: pd.DataFrame,
    register_df: pd.DataFrame,
    docs_dir: Optional[Path],
    input_label: str,
    download_base_name: str,
):
    documents: List[Dict] = st.session_state.get("loaded_documents", [])
    doc_count = len(documents) if documents else len(filtered_df)
    doc_count = max(doc_count, 0)
    has_docs_dir = docs_dir is not None and docs_dir.exists()

    st.markdown(
        "Use this analyser to surface policy pairs that are saying the same thing so you can consider merging or aligning them."
    )
    st.markdown(
        "1. Filter the register (left sidebar) so you’re looking at the right slice.\n"
        "2. Click **Run analyser** below to score overlaps.\n"
        "3. Review the **Consolidation shortlist** to spot the strongest merge opportunities."
    )

    st.markdown("#### Run the analyser")
    max_clusters = max(1, min(15, doc_count if doc_count else 1))
    cluster_count = min(4, max_clusters)
    run_analysis = st.button("Run analyser", key="run_analysis_bert")

    advanced_expander = st.expander("Advanced tuning (optional)", expanded=False)
    with advanced_expander:
        st.caption("Only adjust these if you need to explain a tricky cluster or broaden the cast of policies.")
        adv_col1, adv_col2 = st.columns(2)
        with adv_col1:
            high_df_cutoff = st.slider(
                "Ignore common wording (%)",
                min_value=0,
                max_value=100,
                value=st.session_state.get("high_df_cutoff_slider", 50),
                step=5,
                key="high_df_cutoff_slider",
                help="Filters out boilerplate language that appears in many policies.",
            )
            ontology_weight = st.slider(
                "Blend in shared definitions/legislation",
                min_value=0.0,
                max_value=1.0,
                value=st.session_state.get("ontology_weight_slider", 0.3),
                step=0.05,
                key="ontology_weight_slider",
                help="0 = compare wording only. 1 = compare ontology only. Most teams stay near 0.3.",
            )
        with adv_col2:
            term_bias = st.slider(
                "Lean towards definitions vs legislation",
                min_value=0.0,
                max_value=1.0,
                value=st.session_state.get("ontology_term_bias_slider", 0.5),
                step=0.05,
                key="ontology_term_bias_slider",
                help="0 = emphasise shared legislation; 1 = emphasise shared glossary terms.",
            )

    if run_analysis:
        clear_cached_results()
        documents = []
        st.session_state["load_errors"] = []
        if len(filtered_df) == 0:
            st.info("No documents available. Adjust the register filters to load documents.")
        elif not has_docs_dir:
            st.warning("Set a valid documents directory to load policy documents.")
        else:
            doc_ids = filtered_df["register_id"].astype(str).tolist()
            documents, load_errors = load_documents_from_ids(doc_ids, register_df, docs_dir)
            st.session_state["load_errors"] = load_errors
            if load_errors:
                st.warning("Some documents could not be loaded:")
                preview_errors = load_errors[:5]
                for err in preview_errors:
                    st.write(f"- {err}")
                if len(load_errors) > 5:
                    st.write(f"…and {len(load_errors) - 5} more.")
            if not documents:
                st.warning("No policy documents were loaded.")
            else:
                with st.spinner("Computing document similarities..."):
                    doc_labels: List[str] = []
                    doc_texts: List[str] = []
                    doc_id_list: List[Optional[int]] = []
                    doc_metadata_list: List[Dict] = []
                    for doc in documents:
                        label = doc.get("name") or doc.get("title") or doc.get("source") or "Document"
                        doc_labels.append(label)
                        text = normalize_policy_text(doc.get("raw_text", ""))
                        doc_texts.append(text)
                        try:
                            register_id = int(doc.get("register_id"))
                            doc_id_list.append(register_id)
                        except (TypeError, ValueError):
                            register_id = None
                            doc_id_list.append(None)

                        meta = doc.get("metadata") or {}
                        word_count_value = meta.get("word_count")
                        if isinstance(word_count_value, str):
                            try:
                                word_count_value = int(float(word_count_value.replace(",", "")))
                            except (ValueError, TypeError):
                                word_count_value = None
                        elif isinstance(word_count_value, (float, np.floating)):
                            word_count_value = int(word_count_value)
                        elif not isinstance(word_count_value, (int, np.integer)):
                            word_count_value = None

                        last_revised_value = ensure_timestamp_naive(meta.get("last_revised"))

                        doc_metadata_list.append(
                            {
                                "register_id": register_id,
                                "title": meta.get("display_title") or meta.get("title") or label,
                                "type": meta.get("type"),
                                "policy_owner": meta.get("policy_owner"),
                                "executive_owner": meta.get("executive_owner"),
                                "policy_subject": meta.get("policy_subject"),
                                "policy_sub_topic": meta.get("policy_sub_topic"),
                                "word_count": word_count_value,
                                "last_revised": last_revised_value,
                            }
                        )

                    processed_texts = list(doc_texts)
                    filtered_terms: List[str] = []
                    if high_df_cutoff > 0:
                        processed_texts, filtered_terms = remove_high_frequency_terms(processed_texts, high_df_cutoff)

                    embeddings = compute_text_embeddings(tuple(processed_texts)) if processed_texts else np.empty((0, 0))
                    if embeddings.size == 0:
                        st.error("Unable to compute embeddings for the selected documents.")
                    else:
                        text_similarity_matrix = compute_cosine_similarity_matrix(embeddings)

                        ontology_df = load_ontology_data(str(docs_dir)) if docs_dir else None
                        term_map: Dict[int, Dict[str, str]] = {}
                        legislation_map: Dict[int, Set[str]] = {}
                        if ontology_df is not None and doc_id_list:
                            scoped = ontology_df[
                                ontology_df["register_id"].isin(
                                    [rid for rid in doc_id_list if isinstance(rid, int)]
                                )
                            ].copy()
                            if not scoped.empty:
                                scoped["term"] = scoped["term"].fillna("").astype(str).str.strip()
                                scoped["definition"] = scoped["definition"].fillna("").astype(str).str.strip()
                                scoped["legislation_reference"] = scoped["legislation_reference"].fillna("").astype(str)
                                for rid, group in scoped.groupby("register_id"):
                                    term_map[rid] = {
                                        row.term: normalize_policy_text(row.definition)
                                        for row in group.itertuples()
                                        if row.term
                                    }
                                    legislation_entries: Set[str] = set()
                                    for ref in group["legislation_reference"]:
                                        if not ref:
                                            continue
                                        parts = [part.strip() for part in ref.split(";")]
                                        legislation_entries.update([part for part in parts if part])
                                    legislation_map[rid] = legislation_entries

                        snippets = []
                        for text in doc_texts:
                            snippet = text.strip()
                            if len(snippet) > 400:
                                snippet = snippet[:400] + "…"
                            snippets.append(snippet)

                        effective_clusters = max(1, min(cluster_count, len(documents)))
                        if len(documents) <= 1:
                            cluster_labels = [0] * len(documents)
                            theme_keywords = {0: ["(single document)"]}
                        else:
                            cluster_labels, _ = cluster_segments(embeddings, n_clusters=effective_clusters)
                            theme_keywords = label_clusters_by_tfidf(processed_texts, cluster_labels, n_top=3)

                        hybrid_matrix = text_similarity_matrix.copy()
                        pairwise_details: Dict[Tuple[int, int], Dict] = {}
                        term_weight = max(term_bias, 0.0)
                        legislation_weight = max(1.0 - term_bias, 0.0)
                        if term_weight + legislation_weight == 0:
                            term_weight = legislation_weight = 0.5

                        n_docs = len(documents)
                        for i in range(n_docs):
                            for j in range(i + 1, n_docs):
                                base_score = float(text_similarity_matrix[i, j])
                                rid_i = doc_id_list[i]
                                rid_j = doc_id_list[j]
                                shared_terms: Set[str] = set()
                                conflicts: Set[str] = set()
                                aligned_terms: Set[str] = set()
                                term_union: Set[str] = set()
                                term_score = 0.0
                                conflict_ratio = 0.0
                                if isinstance(rid_i, int) and isinstance(rid_j, int) and term_map:
                                    terms_i = term_map.get(rid_i, {})
                                    terms_j = term_map.get(rid_j, {})
                                    if terms_i or terms_j:
                                        keys_i = set(terms_i.keys())
                                        keys_j = set(terms_j.keys())
                                        shared_terms = keys_i & keys_j
                                        term_union = keys_i | keys_j
                                        if term_union:
                                            term_score = len(shared_terms) / len(term_union)
                                        for term in shared_terms:
                                            def_i = terms_i.get(term, "")
                                            def_j = terms_j.get(term, "")
                                            if def_i and def_j:
                                                if def_i == def_j:
                                                    aligned_terms.add(term)
                                                else:
                                                    conflicts.add(term)
                                        if shared_terms:
                                            conflict_ratio = len(conflicts) / len(shared_terms)

                                shared_legislation: Set[str] = set()
                                legislation_union: Set[str] = set()
                                legislation_score = 0.0
                                if isinstance(rid_i, int) and isinstance(rid_j, int) and legislation_map:
                                    legs_i = legislation_map.get(rid_i, set())
                                    legs_j = legislation_map.get(rid_j, set())
                                    if legs_i or legs_j:
                                        shared_legislation = legs_i & legs_j
                                        legislation_union = legs_i | legs_j
                                        if legislation_union:
                                            legislation_score = len(shared_legislation) / len(legislation_union)

                                ontology_component = 0.0
                                if term_score > 0 or legislation_score > 0:
                                    numerator = term_weight * term_score + legislation_weight * legislation_score
                                    ontology_component = numerator / (term_weight + legislation_weight)

                                hybrid_score = ((1 - ontology_weight) * base_score) + (ontology_weight * ontology_component)
                                hybrid_score = max(0.0, min(1.0, hybrid_score))

                                hybrid_matrix[i, j] = hybrid_matrix[j, i] = hybrid_score

                                pairwise_details[(i, j)] = {
                                    "index_a": i,
                                    "index_b": j,
                                    "label_a": doc_labels[i],
                                    "label_b": doc_labels[j],
                                    "register_id_a": rid_i,
                                    "register_id_b": rid_j,
                                    "hybrid_score": hybrid_score,
                                    "text_score": base_score,
                                    "term_score": term_score,
                                    "legislation_score": legislation_score,
                                    "ontology_component": ontology_component,
                                    "shared_terms": sorted(shared_terms),
                                    "aligned_terms": sorted(aligned_terms),
                                    "conflicting_terms": sorted(conflicts),
                                    "shared_legislation": sorted(shared_legislation),
                                    "all_legislation": sorted(legislation_union),
                                    "all_terms": sorted(term_union),
                                }

                        st.session_state["loaded_documents"] = documents
                        st.session_state["document_embeddings"] = embeddings
                        st.session_state["document_ids"] = doc_id_list
                        st.session_state["document_metadata"] = doc_metadata_list
                        st.session_state["document_texts_tuple"] = tuple(doc_texts)
                        st.session_state["document_similarity_matrix"] = hybrid_matrix
                        st.session_state["text_similarity_matrix"] = text_similarity_matrix
                        st.session_state["document_labels"] = doc_labels
                        st.session_state["document_snippets"] = snippets
                        st.session_state["filtered_terms"] = filtered_terms
                        st.session_state["document_cluster_labels"] = cluster_labels
                        st.session_state["document_theme_keywords"] = theme_keywords
                        st.session_state["pairwise_similarity_details"] = pairwise_details
                        st.session_state["document_terms"] = {
                            rid: dict(terms) for rid, terms in term_map.items()
                        }
                        st.session_state["document_legislation"] = {
                            rid: set(legs) for rid, legs in legislation_map.items()
                        }
                        st.session_state["hybrid_parameters"] = {
                            "ontology_weight": ontology_weight,
                            "term_bias": term_bias,
                        }
                        st.info(f"Overlaps scored for {len(documents)} documents.")

    doc_labels = st.session_state.get("document_labels")
    doc_snippets = st.session_state.get("document_snippets", [])
    removed_terms = st.session_state.get("filtered_terms", [])
    cluster_labels_state = st.session_state.get("document_cluster_labels")
    theme_keywords_state = st.session_state.get("document_theme_keywords", {})
    pairwise_details = st.session_state.get("pairwise_similarity_details", {})
    doc_metadata_state = st.session_state.get("document_metadata", [])
    doc_ids_state = st.session_state.get("document_ids") or []
    doc_texts_tuple = st.session_state.get("document_texts_tuple")
    doc_index_by_id: Dict[int, int] = {
        int(rid): idx
        for idx, rid in enumerate(doc_ids_state)
        if isinstance(rid, (int, np.integer))
    }

    if removed_terms:
        preview = ", ".join(removed_terms[:20])
        if len(removed_terms) > 20:
            preview += f", … (+{len(removed_terms) - 20} more)"
        st.caption(f"Filtered {len(removed_terms)} high-frequency terms: {preview}")

    similarity_matrix = st.session_state.get("document_similarity_matrix")
    if similarity_matrix is None:
        if not has_docs_dir:
            st.info("Set a valid documents directory to load policy documents.")
        elif len(filtered_df) == 0:
            st.info("No documents available. Adjust the register filters to load documents.")
        else:
            st.info("Click **Run analyser** above to compute policy overlaps.")
        return

    if not doc_labels or len(doc_labels) != similarity_matrix.shape[0]:
        doc_labels = [f"Document {i+1}" for i in range(similarity_matrix.shape[0])]

    metadata_by_index: Dict[int, Dict] = {
        idx: doc_metadata_state[idx]
        for idx in range(min(len(doc_metadata_state), len(doc_labels)))
    }

    st.subheader("Consolidation shortlist")
    shortlist_csv_data: Optional[str] = None

    if pairwise_details and doc_labels:
        st.caption("Strongest overlaps first.")

        metadata_by_id: Dict[int, Dict] = {}
        for meta in doc_metadata_state:
            if not isinstance(meta, dict):
                continue
            rid = meta.get("register_id")
            if isinstance(rid, (int, np.integer)):
                metadata_by_id[int(rid)] = meta

        needed_ids: Set[int] = set()
        for detail in pairwise_details.values():
            rid_a = detail.get("register_id_a")
            rid_b = detail.get("register_id_b")
            if isinstance(rid_a, (int, np.integer)):
                needed_ids.add(int(rid_a))
            if isinstance(rid_b, (int, np.integer)):
                needed_ids.add(int(rid_b))

        missing_ids = [rid for rid in needed_ids if rid not in metadata_by_id]
        if missing_ids and "register_id" in register_df.columns:
            register_subset = register_df[register_df["register_id"].isin(missing_ids)].copy()
            for _, row in register_subset.iterrows():
                register_id = int(row["register_id"])
                word_count_val = row.get("word_count")
                if isinstance(word_count_val, str):
                    try:
                        word_count_val = int(float(word_count_val.replace(",", "")))
                    except (ValueError, TypeError):
                        word_count_val = None
                elif isinstance(word_count_val, (float, np.floating)):
                    word_count_val = int(word_count_val)
                elif not isinstance(word_count_val, (int, np.integer)):
                    word_count_val = None

                last_revised_parsed = ensure_timestamp_naive(row.get("last_revised"))

                metadata_by_id[register_id] = {
                    "register_id": register_id,
                    "title": row.get("display_title") or row.get("title"),
                    "type": row.get("type"),
                    "policy_owner": row.get("policy_owner"),
                    "executive_owner": row.get("executive_owner"),
                    "policy_subject": row.get("policy_subject"),
                    "policy_sub_topic": row.get("policy_sub_topic"),
                    "word_count": word_count_val,
                    "last_revised": last_revised_parsed,
                }

        pair_detail_map: Dict[Tuple[int, int], Dict] = {}
        for detail in pairwise_details.values():
            rid_a = detail.get("register_id_a")
            rid_b = detail.get("register_id_b")
            if isinstance(rid_a, (int, np.integer)) and isinstance(rid_b, (int, np.integer)):
                key = tuple(sorted((int(rid_a), int(rid_b))))
                pair_detail_map[key] = detail

        pair_priority_entries: List[Dict] = []
        for pair_key, detail in pair_detail_map.items():
            rid_a, rid_b = pair_key
            meta_a = metadata_by_id.get(rid_a, {})
            meta_b = metadata_by_id.get(rid_b, {})
            subject_a_raw = normalize_text_field(meta_a.get("policy_subject"))
            subject_b_raw = normalize_text_field(meta_b.get("policy_subject"))
            subtopic_a_raw = normalize_text_field(meta_a.get("policy_sub_topic"))
            subtopic_b_raw = normalize_text_field(meta_b.get("policy_sub_topic"))
            subject_match = (
                subject_a_raw
                and subject_b_raw
                and subject_a_raw.casefold() == subject_b_raw.casefold()
            )
            subtopic_match = (
                subtopic_a_raw
                and subtopic_b_raw
                and subtopic_a_raw.casefold() == subtopic_b_raw.casefold()
            )
            bonus = 0.0
            if subject_match:
                bonus += 0.05
            if subtopic_match:
                bonus += 0.05
            weighted_score = min(1.0, float(detail.get("hybrid_score", 0.0)) + bonus)
            pair_priority_entries.append(
                {
                    "pair_key": pair_key,
                    "detail": detail,
                    "score": weighted_score,
                    "meta_a": meta_a,
                    "meta_b": meta_b,
                    "subject_a": subject_a_raw,
                    "subject_b": subject_b_raw,
                    "subtopic_a": subtopic_a_raw,
                    "subtopic_b": subtopic_b_raw,
                    "subject_match": subject_match,
                    "subtopic_match": subtopic_match,
                }
            )

        shortlist_rows: List[Dict] = []
        today = pd.Timestamp.now(tz=None).normalize()
        prioritized_pairs = sorted(
            pair_priority_entries,
            key=lambda item: item["score"],
            reverse=True,
        )
        max_shortlist = 15
        for entry in prioritized_pairs:
            detail = entry["detail"]
            rid_a, rid_b = entry["pair_key"]
            if not isinstance(rid_a, (int, np.integer)) or not isinstance(rid_b, (int, np.integer)):
                continue
            meta_a = entry["meta_a"] or {}
            meta_b = entry["meta_b"] or {}

            owners = {meta_a.get("policy_owner"), meta_b.get("policy_owner")}
            owners_display = ", ".join(sorted({owner for owner in owners if owner})) or "—"

            word_counts = [
                meta.get("word_count")
                for meta in (meta_a, meta_b)
                if meta and isinstance(meta.get("word_count"), (int, np.integer))
            ]
            total_words = sum(word_counts) if word_counts else np.nan

            revision_dates = [
                ensure_timestamp_naive(meta.get("last_revised"))
                for meta in (meta_a, meta_b)
                if meta
            ]
            revision_dates = [ts for ts in revision_dates if ts is not None]
            if revision_dates:
                oldest_revision = min(revision_dates)
                normalized_oldest = oldest_revision.normalize()
                oldest_display = normalized_oldest.date().isoformat()
                staleness_days = (today - normalized_oldest).days
            else:
                oldest_display = "—"
                staleness_days = np.nan

            subject_a = entry["subject_a"]
            subject_b = entry["subject_b"]
            subtopic_a = entry["subtopic_a"]
            subtopic_b = entry["subtopic_b"]
            subject_display = f"{subject_a or '—'} ↔ {subject_b or '—'}"
            subtopic_display = f"{subtopic_a or '—'} ↔ {subtopic_b or '—'}"
            alignment_parts = []
            if entry["subject_match"]:
                alignment_parts.append("Same subject")
            if entry["subtopic_match"]:
                alignment_parts.append("Same sub-topic")
            alignment_label = " & ".join(alignment_parts) if alignment_parts else "Mixed focus"

            signal_parts: List[str] = []
            shared_terms = detail.get("shared_terms") or []
            if shared_terms:
                term_preview = ", ".join(shared_terms[:3])
                if len(shared_terms) > 3:
                    term_preview += ", …"
                signal_parts.append(f"Terms: {term_preview}")
            shared_legislation = detail.get("shared_legislation") or []
            if shared_legislation:
                leg_preview = ", ".join(shared_legislation[:2])
                if len(shared_legislation) > 2:
                    leg_preview += ", …"
                signal_parts.append(f"Acts: {leg_preview}")
            signals_display = " | ".join(signal_parts) if signal_parts else "—"

            alerts = []
            if detail.get("conflicting_terms"):
                alerts.append("Definition mismatch")

            shortlist_rows.append(
                {
                    "Policies": f"{detail.get('label_a')} ↔ {detail.get('label_b')}",
                    "Hybrid score": detail.get("hybrid_score", 0.0),
                    "Shared signals": signals_display,
                    "Total words": total_words,
                    "Owners involved": owners_display,
                    "Subjects": subject_display,
                    "Sub-topics": subtopic_display,
                    "Subject alignment": alignment_label,
                    "Oldest revision": oldest_display,
                    "Staleness (days)": staleness_days,
                    "Alerts": ", ".join(alerts) if alerts else "—",
                    "register_id_a": int(rid_a),
                    "register_id_b": int(rid_b),
                    "pair_key": entry["pair_key"],
                }
            )
            if len(shortlist_rows) >= max_shortlist:
                break

        if shortlist_rows:
            shortlist_df = pd.DataFrame(shortlist_rows)

            total_pairs = len(shortlist_rows)
            st.metric("Pairs surfaced", total_pairs)

            display_df = shortlist_df.drop(
                columns=[
                    col
                    for col in ["register_id_a", "register_id_b", "pair_key", "Total words"]
                    if col in shortlist_df.columns
                ]
            ).rename(
                columns={
                    "Policies": "Policy pair",
                    "Hybrid score": "Match strength",
                    "Shared signals": "Shared signals",
                    "Owners involved": "Owners",
                    "Subjects": "Subjects",
                    "Sub-topics": "Sub-topics",
                    "Subject alignment": "Subject alignment",
                    "Oldest revision": "Oldest update",
                    "Staleness (days)": "Days since update",
                    "Alerts": "Alerts",
                }
            )
            display_df["Match strength"] = display_df["Match strength"].apply(lambda val: round(val, 3))
            display_df["Days since update"] = display_df["Days since update"].apply(
                lambda val: int(val) if pd.notna(val) else pd.NA
            )
            st.dataframe(display_df, hide_index=True, width="stretch")

            st.markdown("#### Deep dive")
            st.caption("Pick a pair to see why it scored highly.")
            label_lookup: Dict[str, Tuple[int, int]] = {
                row["Policies"]: row["pair_key"] for row in shortlist_rows
            }
            select_options = ["(Select pair)"] + list(label_lookup.keys())
            selected_pair_label = st.selectbox(
                "Inspect pair",
                options=select_options,
                key="hybrid_pair_select",
            )
            if selected_pair_label != "(Select pair)":
                pair_key = label_lookup.get(selected_pair_label)
                detail = pair_detail_map.get(pair_key) if pair_key else None
                if detail:
                    st.caption(
                        f"Match strength blends text similarity ({round(detail['text_score'], 3)}) with ontology overlap "
                        f"(terms: {round(detail['term_score'], 3)}, legislation: {round(detail['legislation_score'], 3)})."
                    )
                    rid_a = detail.get("register_id_a")
                    rid_b = detail.get("register_id_b")
                    label_a = detail.get("label_a") or selected_pair_label.split(" ↔ ")[0]
                    label_b = detail.get("label_b") or selected_pair_label.split(" ↔ ")[-1]
                    idx_a = doc_index_by_id.get(int(rid_a)) if isinstance(rid_a, (int, np.integer)) else None
                    idx_b = doc_index_by_id.get(int(rid_b)) if isinstance(rid_b, (int, np.integer)) else None
                    matches_a: List[Tuple[str, str, float]] = []
                    matches_b: List[Tuple[str, str, float]] = []
                    unique_terms_a: Set[str] = set()
                    unique_terms_b: Set[str] = set()
                    has_snippet_source = (
                        doc_texts_tuple
                        and idx_a is not None
                        and idx_b is not None
                        and idx_a < len(doc_texts_tuple)
                        and idx_b < len(doc_texts_tuple)
                    )
                    if has_snippet_source:
                        matches_a, matches_b, divergence_a, divergence_b = extract_top_matching_snippets(
                            doc_texts_tuple[idx_a],
                            doc_texts_tuple[idx_b],
                            top_n=3,
                        )
                        for snippet_text, counterpart_text, _ in matches_a:
                            tokens_a = {
                                token
                                for token in re.findall(r"\w+", snippet_text.lower())
                                if len(token) >= 4
                            }
                            tokens_b = {
                                token
                                for token in re.findall(r"\w+", counterpart_text.lower())
                                if len(token) >= 4
                            }
                            unique_terms_a.update(tokens_a - tokens_b)
                            unique_terms_b.update(tokens_b - tokens_a)
                        for snippet_text, counterpart_text, _ in matches_b:
                            tokens_b = {
                                token
                                for token in re.findall(r"\w+", snippet_text.lower())
                                if len(token) >= 4
                            }
                            tokens_a = {
                                token
                                for token in re.findall(r"\w+", counterpart_text.lower())
                                if len(token) >= 4
                            }
                            unique_terms_b.update(tokens_b - tokens_a)
                            unique_terms_a.update(tokens_a - tokens_b)
                        for snippet_text, counterpart_text, _ in divergence_a:
                            tokens_a = {
                                token
                                for token in re.findall(r"\w+", snippet_text.lower())
                                if len(token) >= 4
                            }
                            tokens_b = {
                                token
                                for token in re.findall(r"\w+", counterpart_text.lower())
                                if len(token) >= 4
                            }
                            unique_terms_a.update(tokens_a - tokens_b)
                            unique_terms_b.update(tokens_b - tokens_a)
                        for snippet_text, counterpart_text, _ in divergence_b:
                            tokens_b = {
                                token
                                for token in re.findall(r"\w+", snippet_text.lower())
                                if len(token) >= 4
                            }
                            tokens_a = {
                                token
                                for token in re.findall(r"\w+", counterpart_text.lower())
                                if len(token) >= 4
                            }
                            unique_terms_b.update(tokens_b - tokens_a)
                            unique_terms_a.update(tokens_a - tokens_b)
                    else:
                        st.info("Snippet preview unavailable for this pair.")

                    summary_text = build_pair_summary(
                        label_a=label_a,
                        label_b=label_b,
                        detail=detail,
                        subject_a=subject_a,
                        subject_b=subject_b,
                        subtopic_a=subtopic_a,
                        subtopic_b=subtopic_b,
                        subject_match=entry["subject_match"],
                        subtopic_match=entry["subtopic_match"],
                        matches_a=matches_a,
                        matches_b=matches_b,
                        unique_tokens_a=unique_terms_a,
                        unique_tokens_b=unique_terms_b,
                    )
                    if summary_text:
                        st.markdown(f"**Summary:** {summary_text}")

                    if matches_a or divergence_a:
                        ensure_comparison_styles()
                        compare_cols = st.columns(2)
                        with compare_cols[0]:
                            st.markdown("**Closest matches**")
                            if matches_a:
                                for idx, (snippet_text, counterpart_text, score) in enumerate(matches_a[:3], start=1):
                                    tokens_a = {
                                        token
                                        for token in re.findall(r"\w+", snippet_text.lower())
                                        if len(token) >= 4
                                    }
                                    tokens_b = {
                                        token
                                        for token in re.findall(r"\w+", counterpart_text.lower())
                                        if len(token) >= 4
                                    }
                                    overlap_tokens = tokens_a & tokens_b
                                    unique_a_tokens = tokens_a - tokens_b
                                    unique_b_tokens = tokens_b - tokens_a
                                    highlighted_a = build_highlighted_text(
                                        snippet_text,
                                        overlap_tokens,
                                        unique_a_tokens,
                                        "unique-a",
                                    )
                                    highlighted_b = build_highlighted_text(
                                        counterpart_text,
                                        overlap_tokens,
                                        unique_b_tokens,
                                        "unique-b",
                                    )
                                    st.markdown(
                                        f'<div class="compare-block">'
                                        f'<div class="compare-title">{label_a} · Score {score:.2f}</div>'
                                        f'<div>{highlighted_a}</div>'
                                        f'<div class="compare-title">{label_b}</div>'
                                        f'<div>{highlighted_b}</div>'
                                        f'</div>',
                                        unsafe_allow_html=True,
                                    )
                            else:
                                st.write("No close matches available.")
                        with compare_cols[1]:
                            st.markdown("**Largest divergences**")
                            if divergence_a or divergence_b:
                                for snippet_text, counterpart_text, score in divergence_a[:2]:
                                    tokens_a = {
                                        token
                                        for token in re.findall(r"\w+", snippet_text.lower())
                                        if len(token) >= 4
                                    }
                                    tokens_b = {
                                        token
                                        for token in re.findall(r"\w+", counterpart_text.lower())
                                        if len(token) >= 4
                                    }
                                    overlap_tokens = tokens_a & tokens_b
                                    unique_a_tokens = tokens_a - tokens_b
                                    unique_b_tokens = tokens_b - tokens_a
                                    highlighted_a = build_highlighted_text(
                                        snippet_text,
                                        overlap_tokens,
                                        unique_a_tokens,
                                        "unique-a",
                                    )
                                    highlighted_b = build_highlighted_text(
                                        counterpart_text,
                                        overlap_tokens,
                                        unique_b_tokens,
                                        "unique-b",
                                    )
                                    st.markdown(
                                        f'<div class="compare-block">'
                                        f'<div class="compare-title">{label_a} · Low similarity {score:.2f}</div>'
                                        f'<div>{highlighted_a}</div>'
                                        f'<div class="compare-title">{label_b}</div>'
                                        f'<div>{highlighted_b}</div>'
                                        f'</div>',
                                        unsafe_allow_html=True,
                                    )
                                for snippet_text, counterpart_text, score in divergence_b[:2]:
                                    tokens_b = {
                                        token
                                        for token in re.findall(r"\w+", snippet_text.lower())
                                        if len(token) >= 4
                                    }
                                    tokens_a = {
                                        token
                                        for token in re.findall(r"\w+", counterpart_text.lower())
                                        if len(token) >= 4
                                    }
                                    overlap_tokens = tokens_a & tokens_b
                                    unique_b_tokens = tokens_b - tokens_a
                                    unique_a_tokens = tokens_a - tokens_b
                                    highlighted_b = build_highlighted_text(
                                        snippet_text,
                                        overlap_tokens,
                                        unique_b_tokens,
                                        "unique-b",
                                    )
                                    highlighted_a = build_highlighted_text(
                                        counterpart_text,
                                        overlap_tokens,
                                        unique_a_tokens,
                                        "unique-a",
                                    )
                                    st.markdown(
                                        f'<div class="compare-block">'
                                        f'<div class="compare-title">{label_b} · Low similarity {score:.2f}</div>'
                                        f'<div>{highlighted_b}</div>'
                                        f'<div class="compare-title">{label_a}</div>'
                                        f'<div>{highlighted_a}</div>'
                                        f'</div>',
                                        unsafe_allow_html=True,
                                    )
                            else:
                                st.write("No clear divergences detected.")
                    elif has_snippet_source:
                        st.info("Matched passages not available for this pair.")

                    col_terms, col_legislation = st.columns(2)
                    with col_terms:
                        st.markdown("**Shared terms**")
                        st.write(", ".join(detail["shared_terms"]) if detail["shared_terms"] else "—")
                        st.markdown("**Conflicting definitions**")
                        st.write(", ".join(detail["conflicting_terms"]) if detail["conflicting_terms"] else "—")
                    with col_legislation:
                        st.markdown("**Shared legislation**")
                        st.write(", ".join(detail["shared_legislation"]) if detail["shared_legislation"] else "—")
                        st.markdown("**All legislation cited**")
                        st.write(", ".join(detail["all_legislation"]) if detail["all_legislation"] else "—")
                    if detail["conflicting_terms"]:
                        st.warning(
                            "These policies define some terms differently. Schedule a quick review to decide which definition should lead."
                        )
                    elif detail["shared_terms"] or detail["shared_legislation"]:
                        st.info(
                            "These policies share important concepts. Consider consolidating content or confirming the responsible owner."
                        )
            shortlist_csv_data = shortlist_df.drop(columns=["pair_key"]).to_csv(index=False)
        else:
            st.info("No high-similarity pairs available for the current selection.")
    else:
        st.info("Run the radar to populate the shortlist of likely consolidation candidates.")

    st.subheader("Downloads")
    if shortlist_csv_data:
        st.download_button(
            "Download shortlist (CSV)",
            data=shortlist_csv_data,
            file_name="consolidation_shortlist.csv",
            key="download_consolidation_shortlist",
        )
    if doc_labels and doc_snippets:
        doc_df = pd.DataFrame({"Document": doc_labels, "Snippet": doc_snippets})
        st.download_button(
            "Download document snippets",
            data=doc_df.to_csv(index=False),
            file_name="document_snippets.csv",
            key="download_doc_snippets",
        )



def extract_documents_from_upload(uploaded_file):
    documents = []
    issues = []
    filename = uploaded_file.name
    lower_name = filename.lower()

    def build_doc_record(name, raw_text, paragraphs):
        display_name = os.path.splitext(os.path.basename(name))[0] or name
        return {
            "name": display_name,
            "raw_text": raw_text,
            "paragraphs": paragraphs,
            "source": filename
        }

    def handle_docx_bytes(doc_bytes, label):
        size = len(doc_bytes)
        if size < 1024 and not doc_bytes.startswith(b"PK"):
            issues.append(f"{label}: very small file ({size} bytes) that is not a DOCX package; it may be a shortcut or legacy .doc file.")
            return None
        try:
            doc = Document(BytesIO(doc_bytes))
        except (PackageNotFoundError, BadZipFile, ValueError) as exc:
            msg = str(exc)
            if "File is not a zip file" in msg:
                issues.append(f"{label}: not a DOCX package; convert the original document to .docx before uploading.")
            else:
                issues.append(f"{label}: {msg}")
            return None
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        raw_text = "\n".join(p.text.strip() for p in paragraphs)
        return build_doc_record(label, raw_text, paragraphs)

    try:
        if lower_name.endswith(".zip"):
            file_bytes = uploaded_file.read()
            try:
                with ZipFile(BytesIO(file_bytes)) as zipped:
                    for inner_name in sorted(zipped.namelist()):
                        if inner_name.endswith("/") or "__MACOSX" in inner_name:
                            continue
                        base_name = os.path.basename(inner_name)
                        if base_name.startswith("._"):
                            continue
                        inner_lower = inner_name.lower()
                        if inner_lower.endswith(".docx"):
                            try:
                                with zipped.open(inner_name) as inner_file:
                                    doc_bytes = inner_file.read()
                            except Exception as exc:
                                issues.append(f"{inner_name}: {exc}")
                                continue
                            record = handle_docx_bytes(doc_bytes, inner_name)
                            if record:
                                documents.append(record)
                        elif inner_lower.endswith(".txt"):
                            try:
                                with zipped.open(inner_name) as inner_file:
                                    text = inner_file.read().decode("utf-8", errors="ignore")
                            except Exception as exc:
                                issues.append(f"{inner_name}: {exc}")
                                continue
                            paragraphs = [line.strip() for line in text.split("\n") if line.strip()]
                            raw_text = "\n".join(paragraphs)
                            documents.append(build_doc_record(inner_name, raw_text, paragraphs))
            except BadZipFile:
                issues.append(f"Could not read zip archive '{filename}'. The file may be corrupted.")
        elif lower_name.endswith(".docx"):
            file_bytes = uploaded_file.read()
            record = handle_docx_bytes(file_bytes, filename)
            if record:
                documents.append(record)
        elif lower_name.endswith(".txt"):
            text = uploaded_file.read().decode("utf-8", errors="ignore")
            paragraphs = [line.strip() for line in text.split("\n") if line.strip()]
            raw_text = "\n".join(paragraphs)
            documents.append(build_doc_record(filename, raw_text, paragraphs))
    finally:
        try:
            uploaded_file.seek(0)
        except Exception:
            pass

    return documents, issues


st.set_page_config(
    page_title="Victoria Police policy cleanup tool",
    layout="wide"
)
sidebar_width_css = """
    <style>
        section[data-testid="stSidebar"] {
            width: 380px !important;
            flex: 0 0 380px;
        }
        section[data-testid="stSidebar"] > div:first-child {
            width: 380px;
        }
        section[data-testid="stSidebar"] div[data-testid="stSlider"] {
            width: 340px !important;
        }
        section[data-testid="stSidebar"] div[data-testid="stSlider"] > div:first-child {
            width: 340px !important;
        }
    </style>
"""
st.markdown(sidebar_width_css, unsafe_allow_html=True)
st.title("Victoria Police policy analyser")

if "similarity_matrix" not in st.session_state:
    st.session_state["similarity_matrix"] = None
if "original_segments" not in st.session_state:
    st.session_state["original_segments"] = []
if "cleaned_segments" not in st.session_state:
    st.session_state["cleaned_segments"] = []
if "segment_ranges" not in st.session_state:
    st.session_state["segment_ranges"] = []
if "preview_snippets" not in st.session_state:
    st.session_state["preview_snippets"] = []

sidebar = st.sidebar
sidebar.markdown("### Policy library filters")

if "preprocessing_opts" not in st.session_state:
    st.session_state["preprocessing_opts"] = {}
docs_dir_input = guess_default_docs_dir()
st.session_state["docs_dir_input"] = docs_dir_input

docs_dir: Optional[Path] = Path(docs_dir_input).resolve() if docs_dir_input else None
register_df: Optional[pd.DataFrame] = None
register_error: Optional[str] = None

if docs_dir_input and docs_dir.exists():
    register_path = docs_dir / DEFAULT_REGISTER_FILENAME
    if register_path.exists():
        try:
            register_df = load_register_data(str(register_path))
        except Exception as exc:  # pylint: disable=broad-except
            register_error = f"Could not read {DEFAULT_REGISTER_FILENAME}: {exc}"
    else:
        register_error = f"Could not find '{DEFAULT_REGISTER_FILENAME}' in {docs_dir}"
else:
    register_error = "Could not resolve the packaged data directory."

if register_df is None:
    st.session_state.pop("register_df", None)
    st.error(register_error)
    st.stop()

st.session_state["register_df"] = register_df

if ("register_embeddings" not in st.session_state
        or st.session_state.get("register_embeddings_rows") != len(register_df)):
    embeddings, register_ids = load_or_build_register_embeddings(register_df, docs_dir)
    st.session_state["register_embeddings"] = embeddings
    st.session_state["register_embedding_ids"] = register_ids.tolist()
    st.session_state["register_embeddings_rows"] = len(register_df)

open_target = st.query_params.get("open_file")
if open_target:
    if isinstance(open_target, (list, tuple)) and open_target:
        raw_target = open_target[0]
    else:
        raw_target = open_target
    decoded_path = Path(unquote(raw_target)).resolve() if raw_target else None
    docs_root = docs_dir.resolve() if docs_dir else None
    success = False
    if decoded_path and docs_root and decoded_path.exists() and decoded_path.is_file() and docs_root in decoded_path.parents:
        try:
            if sys.platform.startswith('win'):
                os.startfile(str(decoded_path))  # type: ignore[attr-defined]
            elif sys.platform == 'darwin':
                subprocess.run(['open', str(decoded_path)], check=False)
            else:
                subprocess.run(['xdg-open', str(decoded_path)], check=False)
            success = True
        except Exception as exc:  # pylint: disable=broad-except
            st.warning(f'Could not open document: {exc}')
    if not success:
        st.warning('Unable to open the requested document.')
    try:
        del st.query_params["open_file"]
    except KeyError:
        st.query_params.clear()


type_options = ["All"] + sorted(register_df["type"].dropna().unique()) if "type" in register_df.columns else ["All"]
default_index = type_options.index("VPM") if "VPM" in type_options else 0
type_filter = sidebar.selectbox(
    "3. Show policy type",
    options=type_options,
    index=default_index,
    key="register_type_filter"
)

filtered_df = register_df.copy()
if type_filter != "All" and "type" in filtered_df.columns:
    filtered_df = filtered_df[filtered_df["type"] == type_filter]

filtered_df = filtered_df.sort_values(by=["display_title", "register_id"])

display_columns = ["register_id", "display_title", "type", "last_revised", "word_count", "relative_path", "policy_owner"]
available_columns = [col for col in display_columns if col in filtered_df.columns]
register_table = filtered_df[available_columns].copy()
if docs_dir and not register_table.empty:
    register_table["file_link"] = register_table["relative_path"].apply(
        lambda rp: (docs_dir / rp).resolve().as_uri() if rp else ""
    )
else:
    register_table["file_link"] = ""
if "last_revised" in register_table.columns:
    register_table["last_revised"] = register_table["last_revised"].apply(
        lambda x: x.date().isoformat() if isinstance(x, pd.Timestamp) and not pd.isna(x) else ""
    )
rename_map = {
    "register_id": "ID",
    "display_title": "Title",
    "type": "Type",
    "last_revised": "Last revised",
    "word_count": "Word count",
    "relative_path": "Path",
    "policy_owner": "Policy owner",
    "file_link": "Document",
}
register_table = register_table.rename(columns=rename_map)
ordered_cols = [col for col in ["ID", "Title", "Type", "Last revised", "Word count", "Policy owner", "Document"] if col in register_table.columns]
if ordered_cols:
    register_table = register_table[ordered_cols]
link_column_config = {}
if "Document" in register_table.columns:
    link_column_config["Document"] = st.column_config.LinkColumn("Document", display_text="Open document")

selection_signature = (type_filter, docs_dir_input)
if selection_signature != st.session_state.get("selection_signature"):
    clear_cached_results()
    st.session_state["selection_signature"] = selection_signature

load_errors = st.session_state.get("load_errors", [])
if load_errors:
    error_preview = "\n".join(load_errors[:5])
    more_count = max(0, len(load_errors) - 5)
    if more_count:
        error_preview += f"\n...and {more_count} more."
    sidebar.warning(f"Issues loading selected documents:\n{error_preview}")

doc_count = len(filtered_df)
if type_filter != "All":
    input_label = f"{doc_count} register documents ({type_filter})"
else:
    input_label = f"{doc_count} register documents"

download_base_name = re.sub(r"[^A-Za-z0-9_-]+", "_", input_label).strip("_") or "policy_similarity"

tabs = st.tabs(
    [
        "Policy library",
        "Policy consolidation analyser",
        "Legislation matcher",
    ]
)

explorer_tab, doc_tab, legislation_tab = tabs

with doc_tab:
    st.subheader("Policy consolidation analyser")
    render_document_similarity_tab(
        filtered_df=filtered_df,
        register_df=register_df,
        docs_dir=docs_dir,
        input_label=input_label,
        download_base_name=download_base_name,
    )

with explorer_tab:
    st.subheader("Policy library")
    st.markdown(
        "Search or filter the official policy register. Type a keyword or question into the search bar to search both titles and body text."
    )
    st.caption(f"{doc_count} policies match the current filters.")

    # Ensure chunk embeddings are available
    chunk_state = st.session_state.get("policy_chunk_state")
    register_signature = tuple(sorted(register_df["register_id"].astype(int).tolist()))
    if (
        chunk_state is None
        or chunk_state.get("register_signature") != register_signature
    ):
        with st.spinner("Preparing search index…"):
            chunk_df, chunk_embeddings = load_policy_chunks(register_df)
        st.session_state["policy_chunk_state"] = {
            "df": chunk_df,
            "embeddings": chunk_embeddings,
            "register_signature": register_signature,
        }

    chunk_state = st.session_state.get("policy_chunk_state", {})
    chunk_df = chunk_state.get("df", pd.DataFrame())
    chunk_embeddings = chunk_state.get("embeddings", np.empty((0, 0)))

    if chunk_df.empty or chunk_embeddings.size == 0:
        st.warning("No policy text available for search. Check that the register includes body text.")
        chunk_df = pd.DataFrame()
        chunk_embeddings = np.empty((0, 0))

    st.session_state.setdefault("policy_keyword_results", [])
    st.session_state.setdefault("policy_keyword_snippets", {})

    refine_terms = []
    selected_subjects: List[str] = []
    selected_subtopics: List[str] = []
    selected_owners: List[str] = []

    left_col, right_col = st.columns([2, 1])

    with left_col:
        st.markdown("#### Keyword or question search")
        search_query = st.text_input(
            "Type a keyword or question",
            st.session_state.get("policy_last_query", ""),
            key="policy_search_input",
        )
        run_search = st.button("Search", key="policy_search_button")

        with st.expander("Refine results", expanded=False):
            refine_terms_input = st.text_input(
                "Refine by keywords (optional, comma separated)",
                "",
                key="policy_search_refine_terms",
            )
            refine_terms = [term.strip() for term in refine_terms_input.split(",") if term.strip()]

            filter_col1, filter_col2, filter_col3 = st.columns(3)
            subject_options = sorted(
                {val for val in chunk_df.get("policy_subject", pd.Series()).dropna().unique() if val}
            )
            subtopic_options = sorted(
                {val for val in chunk_df.get("policy_sub_topic", pd.Series()).dropna().unique() if val}
            )
            owner_options = sorted(
                {val for val in chunk_df.get("policy_owner", pd.Series()).dropna().unique() if val}
            )

            selected_subjects = filter_col1.multiselect(
                "Policy subject",
                options=subject_options,
                key="policy_subject_filter",
            )
            selected_subtopics = filter_col2.multiselect(
                "Policy sub-topic",
                options=subtopic_options,
                key="policy_subtopic_filter",
            )
            selected_owners = filter_col3.multiselect(
                "Policy owner",
                options=owner_options,
                key="policy_owner_filter",
            )

    with right_col:
        st.markdown("#### Recently asked")
        qa_history = st.session_state.get("policy_qa_history", [])
        if qa_history:
            for entry in qa_history[:5]:
                st.caption(entry)
        else:
            st.caption("Run a search to populate this list.")

    query_text = search_query.strip()
    filter_signature = (
        query_text,
        tuple(refine_terms),
        tuple(sorted(selected_subjects)),
        tuple(sorted(selected_subtopics)),
        tuple(sorted(selected_owners)),
        type_filter,
        str(docs_dir) if docs_dir else "",
    )
    last_signature = st.session_state.get("policy_last_signature")
    has_previous_signature = last_signature is not None

    should_run = bool(query_text) and (run_search or (has_previous_signature and filter_signature != last_signature))

    if should_run:
        st.session_state["policy_last_signature"] = filter_signature
        st.session_state["policy_last_query"] = query_text
        st.session_state.setdefault("policy_qa_history", [])
        if query_text not in st.session_state["policy_qa_history"]:
            st.session_state["policy_qa_history"].insert(0, query_text)

        active_ids: Set[int] = set()
        if "register_id" in filtered_df.columns and not filtered_df.empty:
            try:
                active_ids = set(filtered_df["register_id"].dropna().astype(int).tolist())
            except ValueError:
                active_ids = set()
        if not active_ids and not chunk_df.empty:
            try:
                active_ids = set(chunk_df["register_id"].dropna().astype(int).tolist())
            except ValueError:
                active_ids = set()

        available_chunks, embeddings_subset = filter_chunk_subset(
            chunk_df=chunk_df,
            chunk_embeddings=chunk_embeddings,
            active_ids=active_ids,
            selected_subjects=selected_subjects,
            selected_subtopics=selected_subtopics,
            selected_owners=selected_owners,
            refine_terms=refine_terms,
        )

        st.session_state["policy_keyword_results"] = []
        st.session_state["policy_keyword_snippets"] = {}
        st.session_state["policy_qa_results"] = []
        st.session_state["policy_qa_question"] = query_text

        if not available_chunks.empty and embeddings_subset.size > 0:
            enhanced_query = query_text
            if refine_terms:
                enhanced_query = f"{query_text} " + " ".join(refine_terms)

            top_chunks = retrieve_relevant_chunks(
                query=enhanced_query,
                chunk_df=available_chunks,
                chunk_embeddings=embeddings_subset,
                top_k=80,
            )

            if not top_chunks.empty:
                top_chunk_indices = top_chunks.index.to_numpy(dtype=int)
                top_embeddings = embeddings_subset[top_chunk_indices]
                top_chunks = top_chunks.reset_index(drop=True)
                top_chunks["rank"] = top_chunks["score"].rank(method="first", ascending=False)

                grouped = (
                    top_chunks.groupby("register_id")
                    .agg(
                        policy_title=("policy_title", "first"),
                        similarity=("score", "max"),
                        type=("type", "first"),
                        policy_subject=("policy_subject", "first"),
                        policy_sub_topic=("policy_sub_topic", "first"),
                        policy_owner=("policy_owner", "first"),
                        policy_word_count=("policy_word_count", "first"),
                        last_revised=("last_revised", "first"),
                        relative_path=("relative_path", "first"),
                        best_rank=("rank", "min"),
                    )
                    .sort_values(by=["similarity", "best_rank"], ascending=[False, True])
                    .head(20)
                    .reset_index()
                )

                display_records = grouped.rename(
                    columns={
                        "register_id": "ID",
                        "policy_title": "Policy",
                        "similarity": "Similarity",
                        "type": "Type",
                        "policy_subject": "Policy subject",
                        "policy_sub_topic": "Policy sub-topic",
                        "policy_owner": "Policy owner",
                        "policy_word_count": "Word count",
                        "last_revised": "Last revised",
                    }
                )

                display_records["Similarity"] = display_records["Similarity"].round(3)
                for column in ["Policy subject", "Policy sub-topic", "Policy owner", "Last revised"]:
                    display_records[column] = display_records[column].fillna("—")

                snippets_map: Dict[int, List[Dict[str, object]]] = {}
                max_snippets_per_policy = 3

                for _, record in display_records.iterrows():
                    rid = record["ID"]
                    policy_chunks = (
                        top_chunks[top_chunks["register_id"] == rid]
                        .sort_values("score", ascending=False)
                        .head(max_snippets_per_policy)
                    )
                    snippet_records: List[Dict[str, object]] = []
                    for _, chunk_row in policy_chunks.iterrows():
                        snippet_text = str(chunk_row.get("chunk_text", "")).strip()
                        if snippet_text:
                            snippet_clean = re.sub(r"\s+", " ", snippet_text)
                            if len(snippet_clean) > 400:
                                snippet_clean = snippet_clean[:400].rstrip() + "…"
                            snippet_records.append(
                                {"snippet": snippet_clean, "score": float(chunk_row.get("score", 0.0))}
                            )
                    snippets_map[rid] = snippet_records

                def build_document_link(path_value: object) -> str:
                    if docs_dir is None:
                        return ""
                    rel_path = str(path_value or "").strip()
                    if not rel_path:
                        return ""
                    doc_path = (docs_dir / rel_path).expanduser().resolve()
                    if doc_path.exists():
                        return f"?open_file={quote(str(doc_path))}"
                    return ""

                display_records["Document"] = display_records["relative_path"].apply(build_document_link)
                display_records = display_records.drop(columns=["relative_path", "best_rank"])

                st.session_state["policy_keyword_results"] = display_records.to_dict(orient="records")
                st.session_state["policy_keyword_snippets"] = snippets_map
                st.session_state["policy_qa_results"] = answer_question_extractive(
                    question=query_text,
                    available_chunks=top_chunks,
                    embeddings_subset=top_embeddings,
                    docs_dir=docs_dir,
                    top_k=5,
                )
            else:
                st.info("No relevant policies found for that query.")
                st.session_state["policy_keyword_results"] = []
                st.session_state["policy_keyword_snippets"] = {}
                st.session_state["policy_qa_results"] = []
    elif not query_text:
        st.session_state.pop("policy_last_signature", None)
        st.session_state["policy_keyword_results"] = []
        st.session_state["policy_keyword_snippets"] = {}
        st.session_state["policy_qa_results"] = []
        st.session_state.pop("policy_qa_question", None)
    keyword_results_records = st.session_state.get("policy_keyword_results")
    keyword_snippets_map = st.session_state.get("policy_keyword_snippets", {})

    qa_results = st.session_state.get("policy_qa_results")
    qa_question = st.session_state.get("policy_qa_question")

    if qa_results and qa_question:
        st.markdown(f"#### Results for: “{qa_question}”")
        for res in qa_results:
            st.markdown(f"**{res['policy']}** — similarity {res['score']:.3f}")
            st.caption(
                f"Subject: {res.get('policy_subject') or '—'} | Owner: {res.get('policy_owner') or '—'}"
            )
            st.write(res["best_sentence"])
            if res["support_sentences"]:
                with st.expander("Supporting context", expanded=False):
                    for sent in res["support_sentences"]:
                        st.write(f"· {sent}")
            if res["document_link"]:
                st.markdown(f"[Open document]({res['document_link']})")
    elif qa_question:
        st.info("No relevant policies found for that search.")

    st.markdown('#### Full register view')
    st.dataframe(
        register_table,
        hide_index=True,
        column_config=link_column_config if link_column_config else None,
    )
    st.caption("This table respects the filters on the left. Hover to see controls to download or copy rows.")

with legislation_tab:
    st.subheader("Legislation matcher")
    st.markdown(
        "Upload a Word copy of legislation to spot the policies that relate to it. "
        "The tool will analyse the Act and match policies against each part."
    )

    if "legislation_cache" not in st.session_state:
        st.session_state["legislation_cache"] = {}

    cached_items = list(st.session_state["legislation_cache"].items())
    cached_options = [("__none__", "(Select cached Act)")] + [
        (key, f"{value['display_name']} ({value['source_name']})") for key, value in cached_items
    ]

    upload_col, load_col = st.columns([3, 1])
    with upload_col:
        uploaded_legislation = st.file_uploader(
            "Upload Act (.docx)",
            type=["docx"],
            help="Choose the original Word document for the Act you want to compare against policy documents.",
            key="legislation_docx_upload",
        )
    with load_col:
        load_uploaded = st.button("Load", key="process_legislation_upload")

    selected_cached = st.selectbox(
        "Reuse an Act you've already processed",
        options=cached_options,
        format_func=lambda option: option[1],
        index=0,
        key="legislation_cached_select",
    )
    sample_available = LEGISLATION_SAMPLE_FILE.exists()
    load_sample = st.button(
        "Load sample Act",
        key="load_sample_legislation",
        disabled=not sample_available,
    )

    if load_uploaded:
        if uploaded_legislation is None:
            st.warning("Upload a Word document first.")
        else:
            try:
                processed = process_legislation_bytes(uploaded_legislation.getvalue(), uploaded_legislation.name)
            except Exception as exc:  # pylint: disable=broad-except
                st.error(f"Could not process the uploaded Act: {exc}")
            else:
                st.session_state["legislation_cache"][processed["cache_key"]] = processed
                st.session_state["legislation_state"] = processed
                st.success(f"Loaded {processed['display_name']} ({processed['section_count']} sections).")

    cache_key = selected_cached[0]
    if cache_key != "__none__":
        cached_value = st.session_state["legislation_cache"].get(cache_key)
        if cached_value is not None:
            st.session_state["legislation_state"] = cached_value
        else:
            st.warning("The selected Act could not be found in the cache.")

    if load_sample and sample_available:
        with st.spinner("Loading Youth Justice Act sample…"):
            df, embeddings, groups = load_legislation_sections(str(LEGISLATION_SAMPLE_FILE))
        if df is None or embeddings is None or groups is None:
            st.error(f"Could not load sample data from {LEGISLATION_SAMPLE_FILE}.")
        else:
            sample_data = {
                "df": df,
                "embeddings": embeddings,
                "groups": groups,
                "display_name": "Youth Justice Act 2024 (sample)",
                "source_name": LEGISLATION_SAMPLE_FILE.name,
                "cache_key": "sample_youth_justice_act_2024",
                "section_count": len(df),
                "aliases": build_legislation_aliases("Youth Justice Act 2024"),
            }
            st.session_state["legislation_cache"][sample_data["cache_key"]] = sample_data
            st.session_state["legislation_state"] = sample_data
            st.success("Youth Justice Act 2024 sample loaded.")

    legislation_state = st.session_state.get("legislation_state")

    if legislation_state is None:
        st.info("Upload an Act or reuse a cached version to begin.")
    else:
        legislation_df = legislation_state["df"]
        legislation_embeddings = legislation_state["embeddings"]
        legislation_groups = legislation_state["groups"]

        st.markdown(
            f"**Current Act:** {legislation_state['display_name']} "
            f"({legislation_state['section_count']} sections)."
        )

        if not ensure_policy_data(filtered_df, register_df, docs_dir):
            st.stop()

        documents = st.session_state.get("loaded_documents", [])
        doc_embeddings = st.session_state.get("document_embeddings")
        doc_ids = st.session_state.get("document_ids")
        doc_texts_tuple = st.session_state.get("document_texts_tuple")

        if doc_embeddings is None or getattr(doc_embeddings, "size", 0) == 0 or not documents or not doc_ids:
            st.info("Unable to prepare policy data for comparison.")
            st.stop()
        else:
            doc_ids_list = doc_ids if doc_ids else [None] * len(documents)

            st.caption(
                "Choose a legislation level to see the policies that most closely align. "
                "Use this to coordinate updates or confirm responsibility."
            )
            view_level = st.radio(
                "Legislation level",
                options=["Section", "Part", "Chapter"],
                horizontal=True,
            )

            if view_level == "Section":
                section_options = legislation_df.index.tolist()
                section_labels = [
                    f"{int(row.segment_number)}: {row.legislation_section}"
                    for _, row in legislation_df.iterrows()
                ]
                selected_idx = st.selectbox(
                    "Legislation section",
                    section_options,
                    format_func=lambda i: section_labels[i],
                )
                selected_row = legislation_df.iloc[selected_idx]
                section_label = selected_row["legislation_section"]
                section_content = selected_row["content"]
                section_numbers = [selected_row["segment_number"]]
                legislation_vector = legislation_embeddings[selected_idx]
            else:
                group_entries = legislation_groups.get(view_level, [])
                if not group_entries:
                    st.info(f"No {view_level.lower()} groupings available in the legislation data.")
                    st.stop()
                option_indices = list(range(len(group_entries)))
                labels = [
                    f"{entry['label']} (covers {len(entry['segment_numbers'])} sections)"
                    for entry in group_entries
                ]
                selected_idx = st.selectbox(
                    f"Legislation {view_level.lower()}",
                    option_indices,
                    format_func=lambda idx: labels[idx],
                )
                selected_entry = group_entries[selected_idx]
                section_label = selected_entry["label"]
                section_content = selected_entry["content"]
                section_numbers = selected_entry["segment_numbers"]
                legislation_vector = compute_text_embeddings((section_content,))[0]

            st.markdown(f"**Context:** {section_label}")
            st.caption(f"Includes legislation sections: {section_numbers}")
            st.text_area("Legislation content", section_content, height=200)

            max_matches = max(1, len(documents))
            default_matches = min(10, max_matches)
            top_n = st.slider(
                "Number of matching policies to review",
                min_value=1,
                max_value=min(50, max_matches),
                value=default_matches,
            )

            ontology_weight = st.session_state.get("ontology_weight_slider", 0.3)
            term_bias = st.session_state.get("ontology_term_bias_slider", 0.5)
            term_map = st.session_state.get("document_terms", {})
            legislation_map = st.session_state.get("document_legislation", {})
            act_aliases = legislation_state.get("aliases") or build_legislation_aliases(legislation_state["display_name"])
            section_text_lower = section_content.lower()

            hybrid_results = compute_hybrid_scores(
                base_vector=legislation_vector,
                doc_embeddings=doc_embeddings,
                doc_ids=doc_ids_list,
                term_map=term_map,
                legislation_map=legislation_map,
                ontology_weight=ontology_weight,
                term_bias=term_bias,
                section_text_lower=section_text_lower,
                act_aliases=set(act_aliases or []),
            )

            if not hybrid_results:
                st.info("No similarity scores available. Run the Policy consolidation analyser again.")
                st.stop()

            top_entries = sorted(hybrid_results, key=lambda item: item["hybrid_score"], reverse=True)[:top_n]

            rows = []
            for entry in top_entries:
                doc_idx = entry["index"]
                doc_record = documents[doc_idx]
                meta = doc_record.get("metadata", {})
                register_id = doc_record.get("register_id")
                title = meta.get("display_title") or doc_record.get("name")
                policy_type = meta.get("type", "")
                policy_owner = meta.get("policy_owner", "")
                last_revised = meta.get("last_revised")
                if isinstance(last_revised, pd.Timestamp):
                    last_revised = last_revised.date().isoformat()
                elif last_revised:
                    last_revised = str(last_revised)
                relative_path = meta.get("relative_path")
                if docs_dir and relative_path:
                    doc_path = (docs_dir / relative_path).resolve()
                    policy_link = f"?open_file={quote(str(doc_path))}" if doc_path.exists() else ""
                else:
                    policy_link = ""
                snippet_source = doc_texts_tuple[doc_idx] if doc_texts_tuple else doc_record.get("raw_text", "")
                snippet = re.sub(r"\s+", " ", snippet_source).strip()
                if len(snippet) > 300:
                    snippet = snippet[:300] + "…"
                signals = []
                if entry["shared_terms"]:
                    preview_terms = ", ".join(entry["shared_terms"][:3])
                    if len(entry["shared_terms"]) > 3:
                        preview_terms += ", …"
                    signals.append(f"Terms: {preview_terms}")
                if entry["shared_legislation"]:
                    signals.append("References this Act")
                overlap_summary = "; ".join(signals) if signals else "—"

                rows.append(
                    {
                        "Policy": title,
                        "Register ID": register_id,
                        "Type": policy_type,
                        "Policy owner": policy_owner,
                        "Last revised": last_revised,
                        "Hybrid score": round(entry["hybrid_score"], 3),
                        "Text score": round(entry["text_score"], 3),
                        "Overlap": overlap_summary,
                        "Snippet": snippet,
                        "Document": policy_link,
                    }
                )

            results_df = pd.DataFrame(rows)
            policy_link_config = {"Document": st.column_config.LinkColumn("Document", display_text="Open document")}
            st.dataframe(results_df, hide_index=True, column_config=policy_link_config)
            st.caption(
                "Review these policies to confirm consistency with the Act. Higher similarity scores point to closer alignment."
            )
