"""Load and chunk documents for the GPT Agent."""
from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Iterator, List, Tuple

import pandas as pd

from .text_utils import chunk_text, clean_text, slugify
from .constants import (
    CATEGORY_DEFINITION,
    CATEGORY_EXCEPTION,
    CATEGORY_RULE,
    CATEGORY_OTHER,
)

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - docx optional at runtime
    Document = None


@dataclass
class SourceDocument:
    doc_id: str
    title: str
    source_type: str
    path: Path
    text: str


@dataclass
class DocumentChunk:
    chunk_id: str
    source_document: SourceDocument
    index: int
    text: str
    category: str = CATEGORY_OTHER
    definition_terms: List[str] = field(default_factory=list)
    score: float | None = None

    def to_metadata(self) -> dict:
        return {
            "chunk_id": self.chunk_id,
            "doc_id": self.source_document.doc_id,
            "title": self.source_document.title,
            "source_type": self.source_document.source_type,
            "path": str(self.source_document.path),
            "index": self.index,
            "text": self.text,
            "category": self.category,
            "definition_terms": self.definition_terms,
        }


DEFINITION_KEYWORDS = ("definition", "defined", "means", "includes")
EXCEPTION_KEYWORDS = ("exception", "exclusion", "limitation", "unless", "contra-indication", "notes", "does not apply")
RULE_KEYWORDS = ("must", "shall", "are required", "is required", "must not", "are prohibited", "will")


def _extract_definition_terms(text: str) -> List[str]:
    patterns = [
        r"(?im)^(?P<term>[A-Za-z][A-Za-z0-9 \-/]{1,60})\s+(?:means|includes)\b",
        r"(?i)\bdefinition of\s+(?P<term>[A-Za-z][A-Za-z0-9 \-/]{1,60})\b",
    ]
    terms: List[str] = []
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            term = match.group("term").strip()
            if term and term.lower() not in {t.lower() for t in terms}:
                terms.append(term)
    return terms


def _classify_chunk(text: str) -> Tuple[str, List[str]]:
    lowered = text.lower()
    definition_terms = _extract_definition_terms(text)
    if definition_terms or any(keyword in lowered[:200] for keyword in DEFINITION_KEYWORDS):
        return CATEGORY_DEFINITION, definition_terms
    if any(keyword in lowered for keyword in EXCEPTION_KEYWORDS):
        return CATEGORY_EXCEPTION, []
    if any(keyword in lowered for keyword in RULE_KEYWORDS):
        return CATEGORY_RULE, []
    return CATEGORY_OTHER, []


def _read_docx(path: Path) -> str:
    if Document is None:
        raise RuntimeError("python-docx is required to read DOCX files")
    document = Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    return clean_text("\n".join(paragraphs))


def _read_text(path: Path) -> str:
    return clean_text(path.read_text(encoding="utf-8", errors="ignore"))


def _read_excel(path: Path) -> str:
    df = pd.read_excel(path, dtype=str)
    text = "\n".join(" ".join(row.dropna().astype(str)) for _, row in df.iterrows())
    return clean_text(text)


def _supported_suffix(path: Path) -> bool:
    return path.suffix.lower() in {".docx", ".txt", ".md", ".rtf", ".csv", ".xlsx"}


def load_source_documents(vpm_dir: Path, legislation_dir: Path, include_legislation: bool = True) -> List[SourceDocument]:
    documents: List[SourceDocument] = []
    for doc_path in sorted(vpm_dir.glob("**/*")):
        if doc_path.is_dir() or doc_path.suffix.lower() != ".docx":
            continue
        text = _read_docx(doc_path)
        doc_id = f"vpm::{slugify(doc_path.stem)}"
        documents.append(SourceDocument(doc_id, doc_path.stem, "VPM", doc_path, text))

    if include_legislation:
        for leg_path in sorted(legislation_dir.glob("**/*")):
            if leg_path.is_dir() or not _supported_suffix(leg_path):
                continue
            suffix = leg_path.suffix.lower()
            if suffix == ".docx":
                text = _read_docx(leg_path)
            elif suffix == ".xlsx":
                text = _read_excel(leg_path)
            else:
                text = _read_text(leg_path)
            doc_id = f"legislation::{slugify(leg_path.stem)}"
            documents.append(SourceDocument(doc_id, leg_path.stem, "LEGISLATION", leg_path, text))
    return documents


def chunk_documents(documents: Iterable[SourceDocument], chunk_size: int, chunk_overlap: int) -> List[DocumentChunk]:
    chunks: List[DocumentChunk] = []
    for doc in documents:
        text_chunks = chunk_text(doc.text, chunk_size, chunk_overlap)
        if not text_chunks:
            continue
        for idx, chunk_value in enumerate(text_chunks):
            chunk_id = f"{doc.doc_id}::chunk-{idx}"
            category, definition_terms = _classify_chunk(chunk_value)
            chunks.append(
                DocumentChunk(
                    chunk_id=chunk_id,
                    source_document=doc,
                    index=idx,
                    text=chunk_value,
                    category=category,
                    definition_terms=definition_terms,
                )
            )
    return chunks


def chunks_to_metadata_jsonl(chunks: List[DocumentChunk], destination: Path) -> None:
    with destination.open("w", encoding="utf-8") as handle:
        for chunk in chunks:
            handle.write(json.dumps(chunk.to_metadata(), ensure_ascii=False) + "\n")


def load_metadata_jsonl(path: Path) -> List[dict]:
    if not path.exists():
        return []
    records: List[dict] = []
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if not line:
                continue
            records.append(json.loads(line))
    return records
